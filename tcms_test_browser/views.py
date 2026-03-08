import csv
import io
import math
from datetime import datetime

from django.contrib.auth.decorators import login_required
from django.db.models import Count, Prefetch, Q
from django.http import HttpResponse, JsonResponse
from django.utils.decorators import method_decorator
from django.views.generic import TemplateView

from tcms.management.models import Product
from tcms.testcases.models import Category, TestCase
from tcms.testplans.models import TestPlan
from tcms.testruns.models import TestExecution, TestRun


# ==========================================
# Report Metadata Helpers
# ==========================================


def _build_report_metadata(request, report_type, products=None, plans=None):
    """Build a metadata dict for report headers."""
    product_id = request.GET.get("product")
    if product_id:
        try:
            product_filter = Product.objects.get(pk=product_id).name
        except Product.DoesNotExist:
            product_filter = "Unknown"
    else:
        product_filter = "All Products"

    return {
        "author": request.user.username,
        "generated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "report_type": report_type,
        "product_filter": product_filter,
        "products": products or [],
        "plans": plans or [],
    }


def _extract_tc_metadata_from_qs(queryset):
    """Extract product and plan metadata from a test case queryset."""
    products = list(
        queryset.values_list("category__product__name", flat=True)
        .distinct()
        .order_by("category__product__name")
    )
    products = [p for p in products if p]

    plans = list(
        queryset.values_list("plan__pk", "plan__name")
        .distinct()
        .order_by("plan__pk")
    )
    plans = [
        "TP-{}: {}".format(pk, name)
        for pk, name in plans if pk is not None
    ]

    return products, plans


def _extract_plan_metadata_from_qs(queryset):
    """Extract product and plan metadata from a test plan queryset."""
    products = list(
        queryset.values_list("product__name", flat=True)
        .distinct()
        .order_by("product__name")
    )
    products = [p for p in products if p]

    plans = list(
        queryset.values_list("pk", "name")
        .distinct()
        .order_by("pk")
    )
    plans = ["TP-{}: {}".format(pk, name) for pk, name in plans]

    return products, plans


def _extract_run_metadata_from_qs(queryset):
    """Extract product and plan metadata from a test run queryset."""
    products = list(
        queryset.values_list("plan__product__name", flat=True)
        .distinct()
        .order_by("plan__product__name")
    )
    products = [p for p in products if p]

    plans = list(
        queryset.values_list("plan__pk", "plan__name")
        .distinct()
        .order_by("plan__pk")
    )
    plans = [
        "TP-{}: {}".format(pk, name)
        for pk, name in plans if pk is not None
    ]

    return products, plans


def _write_csv_metadata(writer, metadata):
    """Write metadata rows to a CSV writer."""
    writer.writerow(["Report", metadata["report_type"]])
    writer.writerow(["Author", metadata["author"]])
    writer.writerow(["Generated", metadata["generated"]])
    writer.writerow(["Product Filter", metadata["product_filter"]])
    if metadata["products"]:
        writer.writerow(["Products", ", ".join(metadata["products"])])
    if metadata["plans"]:
        writer.writerow(["Test Plans", ", ".join(metadata["plans"])])
    writer.writerow([])


def _write_excel_metadata(ws, metadata, start_row=1):
    """Write metadata rows to an Excel worksheet. Returns next available row."""
    from openpyxl.styles import Font

    bold_font = Font(bold=True)
    rows = [
        ("Report", metadata["report_type"]),
        ("Author", metadata["author"]),
        ("Generated", metadata["generated"]),
        ("Product Filter", metadata["product_filter"]),
    ]
    if metadata["products"]:
        rows.append(("Products", ", ".join(metadata["products"])))

    for i, (label, value) in enumerate(rows):
        row = start_row + i
        cell_label = ws.cell(row=row, column=1, value=label)
        cell_label.font = bold_font
        ws.cell(row=row, column=2, value=value)

    next_row = start_row + len(rows)

    if metadata["plans"]:
        cell_label = ws.cell(row=next_row, column=1, value="Test Plans")
        cell_label.font = bold_font
        for plan in metadata["plans"]:
            ws.cell(row=next_row, column=2, value=plan)
            next_row += 1

    return next_row + 1  # blank row after metadata


def _get_excel_styles():
    """Return consistent Excel table styles used across all reports."""
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    return {
        "header_font": Font(bold=True, color="FFFFFF", size=10),
        "header_fill": PatternFill(start_color="0088CE", end_color="0088CE", fill_type="solid"),
        "header_alignment": Alignment(horizontal="center", vertical="center", wrap_text=True),
        "thin_border": Border(
            left=Side(style="thin", color="CCCCCC"),
            right=Side(style="thin", color="CCCCCC"),
            top=Side(style="thin", color="CCCCCC"),
            bottom=Side(style="thin", color="CCCCCC"),
        ),
        "header_border": Border(
            left=Side(style="thin", color="006699"),
            right=Side(style="thin", color="006699"),
            top=Side(style="thin", color="006699"),
            bottom=Side(style="medium", color="006699"),
        ),
        "data_alignment": Alignment(vertical="top", wrap_text=True),
        "stripe_fill": PatternFill(start_color="F2F7FB", end_color="F2F7FB", fill_type="solid"),
    }


def _write_excel_table(ws, headers, rows, start_row=1):
    """Write a styled table to an Excel worksheet. Returns the next available row."""
    styles = _get_excel_styles()

    # Header row
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=start_row, column=col_idx, value=header)
        cell.font = styles["header_font"]
        cell.fill = styles["header_fill"]
        cell.alignment = styles["header_alignment"]
        cell.border = styles["header_border"]
    ws.row_dimensions[start_row].height = 28

    # Data rows
    row_idx = start_row + 1
    for row in rows:
        for col_idx, value in enumerate(row, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = styles["thin_border"]
            cell.alignment = styles["data_alignment"]
            if (row_idx - start_row) % 2 == 0:
                cell.fill = styles["stripe_fill"]
        row_idx += 1

    # Auto-fit column widths
    for col in ws.columns:
        max_length = 0
        col_letter = None
        for cell in col:
            if col_letter is None:
                col_letter = cell.column_letter
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        if col_letter:
            ws.column_dimensions[col_letter].width = min(max(max_length + 3, 10), 50)

    # Freeze header
    ws.freeze_panes = "A{}".format(start_row + 1)

    return row_idx


def _write_docx_table(doc, heading, headers, rows):
    """Write a styled table to a Word document."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    if heading:
        doc.add_heading(heading, level=2)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row styling
    header_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Set cell background to blue
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "0088CE",
        })
        shading.append(shd)

    # Data rows
    for row_num, row in enumerate(rows):
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = row_cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(8)
            # Alternating row backgrounds
            if row_num % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "F2F7FB",
                })
                shading.append(shd)

    doc.add_paragraph()


def _get_pdf_table_style():
    """Return consistent PDF table style used across all reports."""
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle

    return TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0088ce")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("FONTSIZE", (0, 1), (-1, -1), 7),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("BOX", (0, 0), (-1, 0), 1, colors.HexColor("#006699")),
        ("LINEBELOW", (0, 0), (-1, 0), 1.5, colors.HexColor("#006699")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f2f7fb")]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ])


def _write_docx_metadata(doc, metadata):
    """Add metadata paragraph to a Word document."""
    from docx.shared import Pt

    p = doc.add_paragraph()
    items = [
        ("Report: ", metadata["report_type"]),
        ("Author: ", metadata["author"]),
        ("Generated: ", metadata["generated"]),
        ("Product Filter: ", metadata["product_filter"]),
    ]
    if metadata["products"]:
        items.append(("Products: ", ", ".join(metadata["products"])))

    for i, (label, value) in enumerate(items):
        if i > 0:
            run = p.add_run("\n")
            run.font.size = Pt(8)
        run_label = p.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(8)
        run_value = p.add_run(value)
        run_value.font.size = Pt(8)

    if metadata["plans"]:
        run = p.add_run("\n")
        run.font.size = Pt(8)
        run_label = p.add_run("Test Plans:")
        run_label.bold = True
        run_label.font.size = Pt(8)
        for plan in metadata["plans"]:
            run = p.add_run("\n    \u2022 " + plan)
            run.font.size = Pt(8)


def _build_pdf_metadata_elements(metadata, styles):
    """Return a list of PDF elements (Paragraphs + Spacer) for metadata."""
    from reportlab.platypus import Paragraph, Spacer

    meta_style = styles["Normal"]

    lines = [
        "<b>Report:</b> {}".format(metadata["report_type"]),
        "<b>Author:</b> {}".format(metadata["author"]),
        "<b>Generated:</b> {}".format(metadata["generated"]),
        "<b>Product Filter:</b> {}".format(metadata["product_filter"]),
    ]
    if metadata["products"]:
        lines.append(
            "<b>Products:</b> {}".format(", ".join(metadata["products"]))
        )

    elements = []
    for line in lines:
        elements.append(Paragraph(line, meta_style))

    if metadata["plans"]:
        elements.append(Paragraph("<b>Test Plans:</b>", meta_style))
        for plan in metadata["plans"]:
            elements.append(
                Paragraph("&nbsp;&nbsp;&nbsp;&bull; {}".format(plan), meta_style)
            )

    elements.append(Spacer(1, 12))
    return elements


# ==========================================
# Chart Helper Functions
# ==========================================

# Branded chart color palette
CHART_COLORS = [
    "#0088CE",  # blue
    "#3F9C35",  # green
    "#EC7A08",  # orange
    "#CC0000",  # red
    "#703FEC",  # purple
    "#F0AB00",  # yellow
    "#009596",  # teal
    "#5752D1",  # indigo
    "#8BC1F7",  # light blue
    "#BDE2B9",  # light green
]


def _compute_tc_chart_data(queryset):
    """Compute chart data from TestCase queryset."""
    # By status
    by_status = queryset.values("case_status__name").annotate(count=Count("id"))
    status_dict = {item["case_status__name"] or "Unknown": item["count"] for item in by_status}

    # By priority
    by_priority = queryset.values("priority__value").annotate(count=Count("id"))
    priority_dict = {item["priority__value"] or "Unknown": item["count"] for item in by_priority}

    # Automation
    automated_count = queryset.filter(is_automated=True).count()
    manual_count = queryset.filter(is_automated=False).count()
    automation_dict = {}
    if automated_count > 0:
        automation_dict["Automated"] = automated_count
    if manual_count > 0:
        automation_dict["Manual"] = manual_count

    return {
        "by_status": status_dict,
        "by_priority": priority_dict,
        "automation": automation_dict,
    }


def _compute_plan_chart_data(queryset):
    """Compute chart data from TestPlan queryset."""
    # By type
    by_type = queryset.values("type__name").annotate(count=Count("id"))
    type_dict = {item["type__name"] or "Unknown": item["count"] for item in by_type}

    # Active/Inactive
    active_count = queryset.filter(is_active=True).count()
    inactive_count = queryset.filter(is_active=False).count()
    active_dict = {}
    if active_count > 0:
        active_dict["Active"] = active_count
    if inactive_count > 0:
        active_dict["Inactive"] = inactive_count

    return {
        "by_type": type_dict,
        "active_inactive": active_dict,
    }


def _compute_run_chart_data(queryset):
    """Compute chart data from TestRun queryset."""
    # Get all execution statuses for runs in queryset
    run_ids = list(queryset.values_list("id", flat=True))
    exec_by_status = (
        TestExecution.objects.filter(run_id__in=run_ids)
        .values("status__name")
        .annotate(count=Count("id"))
    )
    status_dict = {item["status__name"] or "Unknown": item["count"] for item in exec_by_status}

    # Completion (stopped vs in progress)
    completed_count = queryset.exclude(stop_date=None).count()
    in_progress_count = queryset.filter(stop_date=None).count()
    completion_dict = {}
    if completed_count > 0:
        completion_dict["Completed"] = completed_count
    if in_progress_count > 0:
        completion_dict["In Progress"] = in_progress_count

    return {
        "by_exec_status": status_dict,
        "completion": completion_dict,
    }


def _add_excel_pie_chart(ws, title, data_dict, cell_ref, data_col=20):
    """Add a pie chart to an Excel worksheet with legend below."""
    if not data_dict:
        return

    from openpyxl.chart import PieChart, Reference

    # Write data to hidden area starting at data_col
    labels = list(data_dict.keys())
    values = list(data_dict.values())

    start_row = 1
    for i, (label, value) in enumerate(zip(labels, values)):
        ws.cell(row=start_row + i, column=data_col, value=label)
        ws.cell(row=start_row + i, column=data_col + 1, value=value)

    # Create pie chart
    chart = PieChart()
    chart.title = title
    chart.width = 9
    chart.height = 12
    chart.legend.position = "b"

    # Show percentage + category in data labels
    chart.dataLabels = chart.series and chart.series[0].graphicalProperties or None
    from openpyxl.chart.label import DataLabelList

    chart.dataLabels = DataLabelList()
    chart.dataLabels.showPercent = True
    chart.dataLabels.showCatName = False
    chart.dataLabels.showVal = False

    # Data references
    data = Reference(
        ws,
        min_col=data_col + 1,
        min_row=start_row,
        max_row=start_row + len(values) - 1,
    )
    cats = Reference(
        ws,
        min_col=data_col,
        min_row=start_row,
        max_row=start_row + len(labels) - 1,
    )

    chart.add_data(data, titles_from_data=False)
    chart.set_categories(cats)

    # Add chart to worksheet
    ws.add_chart(chart, cell_ref)


def _create_pdf_pie_drawing(title, data_dict, width=175):
    """Create a reportlab Drawing with a pie chart and legend below."""
    if not data_dict:
        return None

    from reportlab.graphics.charts.piecharts import Pie
    from reportlab.graphics.shapes import Drawing, Rect, String
    from reportlab.lib import colors

    labels = list(data_dict.keys())
    values = list(data_dict.values())
    total = sum(values)
    if total == 0:
        return None

    # Layout
    title_h = 18
    pie_size = width - 40
    pie_top = title_h + 4
    legend_top = pie_top + pie_size + 8
    legend_line_h = 14
    legend_h = len(labels) * legend_line_h + 4
    height = legend_top + legend_h + 4

    drawing = Drawing(width, height)

    # Title centered at top (reportlab y=0 is bottom)
    title_string = String(width / 2, height - title_h + 2, title)
    title_string.textAnchor = "middle"
    title_string.fontSize = 9
    title_string.fontName = "Helvetica-Bold"
    drawing.add(title_string)

    # Pie chart
    pie = Pie()
    pie_y_bottom = height - pie_top - pie_size
    pie.x = (width - pie_size) / 2
    pie.y = pie_y_bottom
    pie.width = pie_size
    pie.height = pie_size
    pie.data = values
    pie.labels = None
    pie.startAngle = 90
    for i in range(len(values)):
        pie.slices[i].fillColor = colors.HexColor(
            CHART_COLORS[i % len(CHART_COLORS)]
        )
        pie.slices[i].strokeColor = colors.white
        pie.slices[i].strokeWidth = 0.5
    drawing.add(pie)

    # Legend below pie
    box_size = 8
    legend_x = 8
    for i, label in enumerate(labels):
        y = height - legend_top - i * legend_line_h - legend_line_h
        color = colors.HexColor(CHART_COLORS[i % len(CHART_COLORS)])
        rect = Rect(legend_x, y + 1, box_size, box_size)
        rect.fillColor = color
        rect.strokeColor = color
        drawing.add(rect)
        pct = round(values[i] / total * 100)
        text = "{} ({}, {}%)".format(label, values[i], pct)
        s = String(legend_x + box_size + 4, y + 1, text)
        s.fontSize = 7
        s.fontName = "Helvetica"
        drawing.add(s)

    return drawing


def _add_pdf_charts_row(chart_items):
    """Create a borderless reportlab Table with charts placed horizontally.

    Returns a Table flowable, or None if no charts produced.
    """
    from reportlab.platypus import Table, TableStyle

    drawings = []
    for title, data_dict in chart_items:
        d = _create_pdf_pie_drawing(title, data_dict)
        if d:
            drawings.append(d)

    if not drawings:
        return None

    table = Table([drawings])
    table.setStyle(
        TableStyle(
            [
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    return table


def _create_chart_image_buffer(title, data_dict, width=260):
    """Create a pie chart PNG buffer for DOCX embedding using Pillow.

    Legend is drawn below the pie chart so the image fits in a narrow
    table column when charts are placed side-by-side.
    """
    if not data_dict:
        return None

    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return None

    labels = list(data_dict.keys())
    values = list(data_dict.values())
    total = sum(values)
    if total == 0:
        return None

    # Parse hex colors to RGB tuples
    rgb_colors = []
    for hex_color in CHART_COLORS:
        h = hex_color.lstrip("#")
        rgb_colors.append(tuple(int(h[i : i + 2], 16) for i in (0, 2, 4)))

    # Try to load a decent font, fall back to default
    try:
        font_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 13)
        font_label = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 10)
    except (IOError, OSError):
        try:
            font_title = ImageFont.truetype("/usr/share/fonts/TTF/DejaVuSans-Bold.ttf", 13)
            font_label = ImageFont.truetype("/usr/share/fonts/TTF/DejaVuSans.ttf", 10)
        except (IOError, OSError):
            font_title = ImageFont.load_default()
            font_label = ImageFont.load_default()

    # Layout constants
    title_top = 6
    title_h = 22
    pie_size = width - 40
    pie_top = title_top + title_h
    legend_top = pie_top + pie_size + 8
    box_size = 10
    line_height = 18
    legend_h = len(labels) * line_height + 6

    height = legend_top + legend_h + 6

    img = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    # Draw title centered
    title_bbox = draw.textbbox((0, 0), title, font=font_title)
    title_w = title_bbox[2] - title_bbox[0]
    draw.text(((width - title_w) // 2, title_top), title, fill=(0, 0, 0), font=font_title)

    # Draw pie chart centered
    pie_x = (width - pie_size) // 2
    pie_box = (pie_x, pie_top, pie_x + pie_size, pie_top + pie_size)

    start_angle = -90
    slice_angles = []
    for value in values:
        sweep = (value / total) * 360
        slice_angles.append((start_angle, start_angle + sweep))
        start_angle += sweep

    for i, (start, end) in enumerate(slice_angles):
        color = rgb_colors[i % len(rgb_colors)]
        draw.pieslice(pie_box, start=start, end=end, fill=color, outline=(255, 255, 255))

    # Draw legend below the pie
    legend_x = 14
    for i, label in enumerate(labels):
        color = rgb_colors[i % len(rgb_colors)]
        y = legend_top + i * line_height
        draw.rectangle(
            [legend_x, y, legend_x + box_size, y + box_size], fill=color
        )
        pct = round(values[i] / total * 100)
        text = "{} ({}, {}%)".format(label, values[i], pct)
        draw.text(
            (legend_x + box_size + 5, y - 1), text, fill=(0, 0, 0), font=font_label
        )

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def _add_docx_charts_row(doc, chart_items):
    """Add chart images in a horizontal borderless table row.

    Args:
        doc: python-docx Document
        chart_items: list of (title, data_dict) tuples
    """
    from docx.shared import Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    buffers = []
    for title, data_dict in chart_items:
        buf = _create_chart_image_buffer(title, data_dict)
        if buf:
            buffers.append(buf)

    if not buffers:
        return

    cols = len(buffers)
    table = doc.add_table(rows=1, cols=cols)

    # Remove all borders
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:{}".format(edge))
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)

    # Insert chart images into cells
    img_width = Inches(2.1) if cols >= 3 else Inches(2.8)
    for i, buf in enumerate(buffers):
        cell = table.cell(0, i)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 1  # CENTER
        run = paragraph.add_run()
        run.add_picture(buf, width=img_width)


def _get_report_queryset(request):
    """Return a filtered queryset based on request GET params."""
    ids = request.GET.get("ids")
    product_id = request.GET.get("product")
    status_id = request.GET.get("status")
    priority_id = request.GET.get("priority")
    is_automated = request.GET.get("is_automated")

    queryset = TestCase.objects.select_related(
        "case_status",
        "priority",
        "author",
        "category",
        "category__product",
    ).order_by("id")

    if ids:
        id_list = [int(x) for x in ids.split(",") if x.strip().isdigit()]
        if id_list:
            return queryset.filter(pk__in=id_list)

    if product_id:
        queryset = queryset.filter(category__product_id=product_id)
    if status_id:
        queryset = queryset.filter(case_status_id=status_id)
    if priority_id:
        queryset = queryset.filter(priority_id=priority_id)
    if is_automated is not None:
        queryset = queryset.filter(is_automated=is_automated == "true")

    return queryset


def _tc_row(tc):
    """Return a list of display values for a single test case."""
    return [
        tc.pk,
        tc.summary,
        tc.category.product.name if tc.category and tc.category.product else "",
        tc.category.name if tc.category else "",
        tc.case_status.name if tc.case_status else "",
        tc.priority.value if tc.priority else "",
        "Yes" if tc.is_automated else "No",
        tc.author.username if tc.author else "",
        tc.create_date.strftime("%Y-%m-%d") if tc.create_date else "",
    ]


REPORT_HEADERS = [
    "ID", "Summary", "Product", "Category",
    "Status", "Priority", "Automated", "Author", "Created",
]


@method_decorator(login_required, name="dispatch")
class LandingPageView(TemplateView):
    """Landing page with links to all browser views."""

    template_name = "tcms_test_browser/landing.html"

    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["total_cases"] = TestCase.objects.count()
        context["total_plans"] = TestPlan.objects.count()
        context["total_runs"] = TestRun.objects.count()
        context["total_executions"] = TestExecution.objects.count()
        return context


@method_decorator(login_required, name="dispatch")
class TestCaseBrowserView(TemplateView):
    """
    Test Case Browser with tree navigation.
    Left panel: Product → Category tree with test case counts
    Right panel: Selected test case details
    """

    template_name = "tcms_test_browser/browser.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        context["active_nav"] = "cases"

        # Get all products with categories and test case counts in 2 queries
        categories_qs = (
            Category.objects.annotate(testcase_count=Count("category_case"))
            .order_by("name")
        )
        products = Product.objects.prefetch_related(
            Prefetch("category", queryset=categories_qs)
        ).order_by("name")

        tree_data = []
        for product in products:
            categories = product.category.all()

            product_testcase_count = sum(c.testcase_count for c in categories)

            category_list = [
                {
                    "id": cat.pk,
                    "name": cat.name,
                    "count": cat.testcase_count,
                }
                for cat in categories
            ]

            tree_data.append(
                {
                    "id": product.pk,
                    "name": product.name,
                    "count": product_testcase_count,
                    "categories": category_list,
                }
            )

        context["tree_data"] = tree_data
        return context


@login_required
def api_testcases_by_category(request, category_id):
    """API endpoint to get test cases for a category."""
    testcases = TestCase.objects.filter(category_id=category_id).select_related(
        "case_status", "priority", "author", "category"
    ).order_by("summary")

    data = [
        {
            "id": tc.pk,
            "summary": tc.summary,
            "case_status": tc.case_status.name if tc.case_status else None,
            "priority": tc.priority.value if tc.priority else None,
            "author": tc.author.username if tc.author else None,
            "is_automated": tc.is_automated,
            "create_date": tc.create_date.isoformat() if tc.create_date else None,
        }
        for tc in testcases
    ]

    return JsonResponse({"testcases": data})


@login_required
def api_testcase_detail(request, testcase_id):
    """API endpoint to get test case details."""
    try:
        tc = TestCase.objects.select_related(
            "case_status", "priority", "author", "default_tester", "reviewer", "category"
        ).get(pk=testcase_id)
    except TestCase.DoesNotExist:
        return JsonResponse({"error": "Test case not found"}, status=404)

    # Get related data
    components = list(tc.component.values_list("name", flat=True))
    tags = list(tc.tag.values_list("name", flat=True))
    plans = list(tc.plan.values("id", "name"))

    data = {
        "id": tc.pk,
        "summary": tc.summary,
        "text": tc.text or "",
        "notes": tc.notes or "",
        "case_status": tc.case_status.name if tc.case_status else None,
        "case_status_id": tc.case_status_id,
        "priority": tc.priority.value if tc.priority else None,
        "priority_id": tc.priority_id,
        "category": tc.category.name if tc.category else None,
        "category_id": tc.category_id,
        "author": tc.author.username if tc.author else None,
        "default_tester": tc.default_tester.username if tc.default_tester else None,
        "reviewer": tc.reviewer.username if tc.reviewer else None,
        "is_automated": tc.is_automated,
        "script": tc.script or "",
        "arguments": tc.arguments or "",
        "requirement": tc.requirement or "",
        "extra_link": tc.extra_link or "",
        "setup_duration": str(tc.setup_duration) if tc.setup_duration else None,
        "testing_duration": str(tc.testing_duration) if tc.testing_duration else None,
        "create_date": tc.create_date.isoformat() if tc.create_date else None,
        "components": components,
        "tags": tags,
        "plans": plans,
    }

    return JsonResponse(data)


@login_required
def api_search_testcases(request):
    """API endpoint to search test cases."""
    query = request.GET.get("q", "").strip()
    product_id = request.GET.get("product")
    category_id = request.GET.get("category")
    status_id = request.GET.get("status")
    is_automated = request.GET.get("is_automated")

    testcases = TestCase.objects.select_related(
        "case_status", "priority", "author", "category", "category__product"
    )

    if query:
        q_filter = Q(summary__icontains=query) | Q(text__icontains=query)
        if query.isdigit():
            q_filter = q_filter | Q(pk=int(query))
        testcases = testcases.filter(q_filter)

    if product_id:
        testcases = testcases.filter(category__product_id=product_id)

    if category_id:
        testcases = testcases.filter(category_id=category_id)

    if status_id:
        testcases = testcases.filter(case_status_id=status_id)

    if is_automated is not None:
        testcases = testcases.filter(is_automated=is_automated == "true")

    testcases = testcases.order_by("summary")[:100]  # Limit results

    data = [
        {
            "id": tc.pk,
            "summary": tc.summary,
            "case_status": tc.case_status.name if tc.case_status else None,
            "priority": tc.priority.value if tc.priority else None,
            "author": tc.author.username if tc.author else None,
            "category": tc.category.name if tc.category else None,
            "product": tc.category.product.name if tc.category and tc.category.product else None,
            "is_automated": tc.is_automated,
        }
        for tc in testcases
    ]

    return JsonResponse({"testcases": data})


@login_required
def api_browse_testcases(request):
    """Paginated browse endpoint for test cases."""
    product_id = request.GET.get("product")
    page = int(request.GET.get("page", 1))
    page_size = int(request.GET.get("page_size", 25))

    queryset = TestCase.objects.select_related(
        "case_status", "category", "category__product",
    ).order_by("-id")

    if product_id:
        queryset = queryset.filter(category__product_id=product_id)

    # Chart drill-down filters
    status = request.GET.get("status")
    if status:
        queryset = queryset.filter(case_status__name=status)

    priority = request.GET.get("priority")
    if priority:
        queryset = queryset.filter(priority__value=priority)

    product_name = request.GET.get("product_name")
    if product_name:
        queryset = queryset.filter(category__product__name=product_name)

    total = queryset.count()
    total_pages = math.ceil(total / page_size) if total > 0 else 1
    page = max(1, min(page, total_pages))
    offset = (page - 1) * page_size

    testcases = queryset[offset:offset + page_size]

    data = [
        {
            "id": tc.pk,
            "summary": tc.summary,
            "product": (
                tc.category.product.name
                if tc.category and tc.category.product
                else None
            ),
            "case_status": tc.case_status.name if tc.case_status else None,
        }
        for tc in testcases
    ]

    return JsonResponse({
        "testcases": data,
        "page": page,
        "total_pages": total_pages,
        "total": total,
    })


@login_required
def api_statistics(request):
    """API endpoint returning test case statistics for charts."""
    product_id = request.GET.get("product")

    queryset = TestCase.objects.all()
    if product_id:
        queryset = queryset.filter(category__product_id=product_id)

    total = queryset.count()

    by_status = list(
        queryset.values("case_status__name")
        .annotate(count=Count("id"))
        .order_by("case_status__name")
    )

    by_priority = list(
        queryset.values("priority__value")
        .annotate(count=Count("id"))
        .order_by("priority__value")
    )

    automated_count = queryset.filter(is_automated=True).count()
    manual_count = total - automated_count

    by_product = list(
        queryset.values("category__product__name")
        .annotate(count=Count("id"))
        .order_by("category__product__name")
    )

    return JsonResponse(
        {
            "total": total,
            "by_status": [
                {"name": item["case_status__name"] or "Unknown", "count": item["count"]}
                for item in by_status
            ],
            "by_priority": [
                {"name": item["priority__value"] or "Unknown", "count": item["count"]}
                for item in by_priority
            ],
            "automation": {"automated": automated_count, "manual": manual_count},
            "by_product": [
                {
                    "name": item["category__product__name"] or "Unknown",
                    "count": item["count"],
                }
                for item in by_product
            ],
        }
    )


@login_required
def api_report(request):
    """API endpoint to download a CSV report of test cases."""
    queryset = _get_report_queryset(request)
    products, plans = _extract_tc_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Cases", products, plans)

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = 'attachment; filename="test_cases_report.csv"'

    writer = csv.writer(response)
    _write_csv_metadata(writer, metadata)
    writer.writerow(REPORT_HEADERS)

    for tc in queryset.iterator():
        writer.writerow(_tc_row(tc))

    return response


@login_required
def api_report_excel(request):
    """API endpoint to download an Excel report of test cases."""
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter

    queryset = _get_report_queryset(request)
    products, plans = _extract_tc_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Cases", products, plans)

    wb = Workbook()

    # Create Overview sheet with charts (first sheet)
    chart_data = _compute_tc_chart_data(queryset)
    if any(chart_data.values()):
        ws_charts = wb.active
        ws_charts.title = "Overview"

        # Add charts horizontally
        _add_excel_pie_chart(ws_charts, "Test Cases by Status", chart_data["by_status"], "A1", data_col=20)
        _add_excel_pie_chart(ws_charts, "Test Cases by Priority", chart_data["by_priority"], "J1", data_col=22)
        if chart_data["automation"]:
            _add_excel_pie_chart(ws_charts, "Automation Status", chart_data["automation"], "S1", data_col=25)

        # Create data sheet
        ws = wb.create_sheet("Test Cases")
    else:
        ws = wb.active
        ws.title = "Test Cases"

    # Metadata rows
    data_start = _write_excel_metadata(ws, metadata)

    # Build data rows
    rows = [_tc_row(tc) for tc in queryset.iterator()]

    # Write styled table
    last_row = _write_excel_table(ws, REPORT_HEADERS, rows, start_row=data_start)

    # Auto-filter on all columns
    if last_row > data_start:
        ws.auto_filter.ref = (
            f"A{data_start}:{get_column_letter(len(REPORT_HEADERS))}{last_row - 1}"
        )

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = 'attachment; filename="test_cases_report.xlsx"'
    return response


@login_required
def api_report_docx(request):
    """API endpoint to download a Word document report of test cases."""
    from docx import Document
    from docx.shared import Inches

    queryset = _get_report_queryset(request)
    products, plans = _extract_tc_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Cases", products, plans)

    doc = Document()
    doc.add_heading("Test Cases Report", level=1)
    _write_docx_metadata(doc, metadata)

    # Add charts if available
    chart_data = _compute_tc_chart_data(queryset)
    chart_items = [
        (t, d)
        for t, d in [
            ("By Status", chart_data.get("by_status")),
            ("By Priority", chart_data.get("by_priority")),
            ("Automation", chart_data.get("automation")),
        ]
        if d
    ]
    if chart_items:
        doc.add_heading("Statistics Overview", level=2)
        _add_docx_charts_row(doc, chart_items)
        doc.add_paragraph()

    # Build data rows
    rows = [_tc_row(tc) for tc in queryset.iterator()]

    # Write styled table
    _write_docx_table(doc, None, REPORT_HEADERS, rows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    response["Content-Disposition"] = 'attachment; filename="test_cases_report.docx"'
    return response


@login_required
def api_report_pdf(request):
    """API endpoint to download a PDF report of test cases."""
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table

    queryset = _get_report_queryset(request)
    products, tc_plans = _extract_tc_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Cases", products, tc_plans)
    styles = getSampleStyleSheet()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(letter), topMargin=0.5 * inch)
    elements = []

    elements.append(Paragraph("Test Cases Report", styles["Title"]))
    elements.append(Spacer(1, 12))
    elements.extend(_build_pdf_metadata_elements(metadata, styles))

    # Add charts if available
    chart_data = _compute_tc_chart_data(queryset)
    chart_items = [
        (t, d)
        for t, d in [
            ("By Status", chart_data.get("by_status")),
            ("By Priority", chart_data.get("by_priority")),
            ("Automation", chart_data.get("automation")),
        ]
        if d
    ]
    if chart_items:
        elements.append(Paragraph("Statistics Overview", styles["Heading2"]))
        elements.append(Spacer(1, 12))
        charts_table = _add_pdf_charts_row(chart_items)
        if charts_table:
            elements.append(charts_table)
        elements.append(Spacer(1, 12))

    # Build table data — wrap long text in Paragraphs for word-wrap
    cell_style = styles["Normal"]
    cell_style.fontSize = 7
    cell_style.leading = 9

    table_data = [REPORT_HEADERS]
    for tc in queryset.iterator():
        row = _tc_row(tc)
        # Wrap summary (col 1) in a Paragraph for word-wrap
        row[1] = Paragraph(str(row[1]), cell_style)
        table_data.append(row)

    col_widths = [
        0.5 * inch,   # ID
        2.8 * inch,   # Summary
        1.0 * inch,   # Product
        1.0 * inch,   # Category
        0.8 * inch,   # Status
        0.6 * inch,   # Priority
        0.7 * inch,   # Automated
        0.8 * inch,   # Author
        0.8 * inch,   # Created
    ]

    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    table.setStyle(_get_pdf_table_style())
    elements.append(table)

    doc.build(elements)
    buf.seek(0)

    response = HttpResponse(buf.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="test_cases_report.pdf"'
    return response


# ==========================================
# Test Plan Browser
# ==========================================

PLAN_REPORT_HEADERS = [
    "ID", "Name", "Product", "Type", "Status",
    "Version", "Author", "Created", "Cases", "Runs",
]


def _get_plan_report_queryset(request):
    """Return a filtered TestPlan queryset based on request GET params."""
    ids = request.GET.get("ids")
    product_id = request.GET.get("product")
    type_id = request.GET.get("type")
    is_active = request.GET.get("is_active")

    queryset = TestPlan.objects.select_related(
        "product",
        "product_version",
        "type",
        "author",
    ).annotate(
        case_count=Count("cases", distinct=True),
        run_count=Count("run", distinct=True),
    ).order_by("id")

    if ids:
        id_list = [int(x) for x in ids.split(",") if x.strip().isdigit()]
        if id_list:
            return queryset.filter(pk__in=id_list)

    if product_id:
        queryset = queryset.filter(product_id=product_id)
    if type_id:
        queryset = queryset.filter(type_id=type_id)
    if is_active is not None:
        queryset = queryset.filter(is_active=is_active == "true")

    return queryset


def _plan_row(plan):
    """Return a list of display values for a single test plan."""
    return [
        plan.pk,
        plan.name,
        plan.product.name if plan.product else "",
        plan.type.name if plan.type else "",
        "Active" if plan.is_active else "Inactive",
        plan.product_version.value if plan.product_version else "",
        plan.author.username if plan.author else "",
        plan.create_date.strftime("%Y-%m-%d") if plan.create_date else "",
        plan.case_count,
        plan.run_count,
    ]


@method_decorator(login_required, name="dispatch")
class TestPlanBrowserView(TemplateView):
    """
    Test Plan Browser with tree navigation.
    Left panel: Product → TestPlan tree
    Right panel: Selected test plan details
    """

    template_name = "tcms_test_browser/plan_browser.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        context["active_nav"] = "plans"

        plans_qs = TestPlan.objects.select_related(
            "type", "product_version", "author",
        ).order_by("name")

        products = Product.objects.prefetch_related(
            Prefetch("plan", queryset=plans_qs)
        ).order_by("name")

        tree_data = []
        for product in products:
            plans = product.plan.all()
            plan_list = [
                {
                    "id": plan.pk,
                    "name": plan.name,
                    "type": plan.type.name if plan.type else "",
                    "is_active": plan.is_active,
                }
                for plan in plans
            ]
            tree_data.append(
                {
                    "id": product.pk,
                    "name": product.name,
                    "count": len(plan_list),
                    "plans": plan_list,
                }
            )

        context["tree_data"] = tree_data
        return context


@login_required
def api_plan_detail(request, plan_id):
    """API endpoint to get test plan details."""
    try:
        plan = TestPlan.objects.select_related(
            "product", "product_version", "type", "author",
        ).get(pk=plan_id)
    except TestPlan.DoesNotExist:
        return JsonResponse({"error": "Test plan not found"}, status=404)

    cases = list(
        plan.cases.values("id", "summary")
        .order_by("id")[:20]
    )
    runs = list(
        plan.run.values("id", "summary")
        .order_by("-id")[:20]
    )

    data = {
        "id": plan.pk,
        "name": plan.name,
        "text": plan.text or "",
        "is_active": plan.is_active,
        "product": plan.product.name if plan.product else None,
        "product_id": plan.product_id,
        "product_version": plan.product_version.value if plan.product_version else None,
        "type": plan.type.name if plan.type else None,
        "author": plan.author.username if plan.author else None,
        "create_date": plan.create_date.isoformat() if plan.create_date else None,
        "extra_link": plan.extra_link if hasattr(plan, "extra_link") and plan.extra_link else "",
        "cases": cases,
        "runs": runs,
        "case_count": plan.cases.count(),
        "run_count": plan.run.count(),
    }

    return JsonResponse(data)


@login_required
def api_search_plans(request):
    """API endpoint to search test plans."""
    query = request.GET.get("q", "").strip()
    product_id = request.GET.get("product")
    type_id = request.GET.get("type")
    is_active = request.GET.get("is_active")

    plans = TestPlan.objects.select_related(
        "product", "type", "author",
    )

    if query:
        q_filter = Q(name__icontains=query) | Q(text__icontains=query)
        if query.isdigit():
            q_filter = q_filter | Q(pk=int(query))
        plans = plans.filter(q_filter)
    if product_id:
        plans = plans.filter(product_id=product_id)
    if type_id:
        plans = plans.filter(type_id=type_id)
    if is_active is not None:
        plans = plans.filter(is_active=is_active == "true")

    plans = plans.order_by("name")[:100]

    data = [
        {
            "id": p.pk,
            "name": p.name,
            "product": p.product.name if p.product else None,
            "type": p.type.name if p.type else None,
            "is_active": p.is_active,
            "author": p.author.username if p.author else None,
        }
        for p in plans
    ]

    return JsonResponse({"plans": data})


@login_required
def api_browse_plans(request):
    """Paginated browse endpoint for test plans."""
    product_id = request.GET.get("product")
    page = int(request.GET.get("page", 1))
    page_size = int(request.GET.get("page_size", 25))

    queryset = TestPlan.objects.select_related("product").order_by("-id")

    if product_id:
        queryset = queryset.filter(product_id=product_id)

    # Chart drill-down filters
    type_name = request.GET.get("type_name")
    if type_name:
        queryset = queryset.filter(type__name=type_name)

    product_name = request.GET.get("product_name")
    if product_name:
        queryset = queryset.filter(product__name=product_name)

    total = queryset.count()
    total_pages = math.ceil(total / page_size) if total > 0 else 1
    page = max(1, min(page, total_pages))
    offset = (page - 1) * page_size

    plans = queryset[offset:offset + page_size]

    data = [
        {
            "id": p.pk,
            "name": p.name,
            "product": p.product.name if p.product else None,
            "is_active": p.is_active,
        }
        for p in plans
    ]

    return JsonResponse({
        "plans": data,
        "page": page,
        "total_pages": total_pages,
        "total": total,
    })


@login_required
def api_plan_statistics(request):
    """API endpoint returning test plan statistics for charts."""
    product_id = request.GET.get("product")

    queryset = TestPlan.objects.all()
    if product_id:
        queryset = queryset.filter(product_id=product_id)

    total = queryset.count()

    by_type = list(
        queryset.values("type__name")
        .annotate(count=Count("id"))
        .order_by("type__name")
    )

    active_count = queryset.filter(is_active=True).count()
    inactive_count = total - active_count

    by_product = list(
        queryset.values("product__name")
        .annotate(count=Count("id"))
        .order_by("product__name")
    )

    return JsonResponse(
        {
            "total": total,
            "by_type": [
                {"name": item["type__name"] or "Unknown", "count": item["count"]}
                for item in by_type
            ],
            "active_inactive": {"active": active_count, "inactive": inactive_count},
            "by_product": [
                {"name": item["product__name"] or "Unknown", "count": item["count"]}
                for item in by_product
            ],
        }
    )


@login_required
def api_plan_report(request):
    """API endpoint to download a CSV report of test plans."""
    queryset = _get_plan_report_queryset(request)
    products, plans_list = _extract_plan_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Plans", products, plans_list)

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = 'attachment; filename="test_plans_report.csv"'

    writer = csv.writer(response)
    _write_csv_metadata(writer, metadata)
    writer.writerow(PLAN_REPORT_HEADERS)

    for plan in queryset.iterator():
        writer.writerow(_plan_row(plan))

    return response


@login_required
def api_plan_report_excel(request):
    """API endpoint to download an Excel report of test plans."""
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter

    queryset = _get_plan_report_queryset(request)
    products, plans_list = _extract_plan_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Plans", products, plans_list)

    wb = Workbook()

    # Create Overview sheet with charts (first sheet)
    chart_data = _compute_plan_chart_data(queryset)
    if any(chart_data.values()):
        ws_charts = wb.active
        ws_charts.title = "Overview"

        # Add charts horizontally
        _add_excel_pie_chart(ws_charts, "Test Plans by Type", chart_data["by_type"], "A1", data_col=20)
        if chart_data["active_inactive"]:
            _add_excel_pie_chart(ws_charts, "Active vs Inactive", chart_data["active_inactive"], "J1", data_col=22)

        # Create data sheet
        ws = wb.create_sheet("Test Plans")
    else:
        ws = wb.active
        ws.title = "Test Plans"

    # Metadata rows
    data_start = _write_excel_metadata(ws, metadata)

    # Build data rows
    rows = [_plan_row(plan) for plan in queryset.iterator()]

    # Write styled table
    last_row = _write_excel_table(ws, PLAN_REPORT_HEADERS, rows, start_row=data_start)

    # Auto-filter on all columns
    if last_row > data_start:
        ws.auto_filter.ref = (
            f"A{data_start}:{get_column_letter(len(PLAN_REPORT_HEADERS))}{last_row - 1}"
        )

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = 'attachment; filename="test_plans_report.xlsx"'
    return response


@login_required
def api_plan_report_docx(request):
    """API endpoint to download a Word document report of test plans."""
    from docx import Document
    from docx.shared import Inches

    queryset = _get_plan_report_queryset(request)
    products, plans_list = _extract_plan_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Plans", products, plans_list)

    doc = Document()
    doc.add_heading("Test Plans Report", level=1)
    _write_docx_metadata(doc, metadata)

    # Add charts if available
    chart_data = _compute_plan_chart_data(queryset)
    chart_items = [
        (t, d)
        for t, d in [
            ("By Type", chart_data.get("by_type")),
            ("Active vs Inactive", chart_data.get("active_inactive")),
        ]
        if d
    ]
    if chart_items:
        doc.add_heading("Statistics Overview", level=2)
        _add_docx_charts_row(doc, chart_items)
        doc.add_paragraph()

    # Build data rows
    rows = [_plan_row(plan) for plan in queryset.iterator()]

    # Write styled table
    _write_docx_table(doc, None, PLAN_REPORT_HEADERS, rows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    response["Content-Disposition"] = 'attachment; filename="test_plans_report.docx"'
    return response


@login_required
def api_plan_report_pdf(request):
    """API endpoint to download a PDF report of test plans."""
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table

    queryset = _get_plan_report_queryset(request)
    products, plans_list = _extract_plan_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Plans", products, plans_list)
    styles = getSampleStyleSheet()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(letter), topMargin=0.5 * inch)
    elements = []

    elements.append(Paragraph("Test Plans Report", styles["Title"]))
    elements.append(Spacer(1, 12))
    elements.extend(_build_pdf_metadata_elements(metadata, styles))

    # Add charts if available
    chart_data = _compute_plan_chart_data(queryset)
    chart_items = [
        (t, d)
        for t, d in [
            ("By Type", chart_data.get("by_type")),
            ("Active vs Inactive", chart_data.get("active_inactive")),
        ]
        if d
    ]
    if chart_items:
        elements.append(Paragraph("Statistics Overview", styles["Heading2"]))
        elements.append(Spacer(1, 12))
        charts_table = _add_pdf_charts_row(chart_items)
        if charts_table:
            elements.append(charts_table)
        elements.append(Spacer(1, 12))

    cell_style = styles["Normal"]
    cell_style.fontSize = 7
    cell_style.leading = 9

    table_data = [PLAN_REPORT_HEADERS]
    for plan in queryset.iterator():
        row = _plan_row(plan)
        row[1] = Paragraph(str(row[1]), cell_style)
        table_data.append(row)

    col_widths = [
        0.5 * inch,   # ID
        2.5 * inch,   # Name
        1.0 * inch,   # Product
        0.8 * inch,   # Type
        0.7 * inch,   # Status
        0.8 * inch,   # Version
        0.8 * inch,   # Author
        0.8 * inch,   # Created
        0.5 * inch,   # Cases
        0.5 * inch,   # Runs
    ]

    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    table.setStyle(_get_pdf_table_style())
    elements.append(table)

    doc.build(elements)
    buf.seek(0)

    response = HttpResponse(buf.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="test_plans_report.pdf"'
    return response


# ==========================================
# Test Run Browser
# ==========================================

RUN_REPORT_HEADERS = [
    "ID", "Summary", "Plan", "Product", "Build",
    "Manager", "Started", "Stopped", "Total Executions", "Passed", "Failed",
]


def _get_run_report_queryset(request):
    """Return a filtered TestRun queryset based on request GET params."""
    ids = request.GET.get("ids")
    product_id = request.GET.get("product")
    plan_id = request.GET.get("plan")
    build_id = request.GET.get("build")

    queryset = TestRun.objects.select_related(
        "plan",
        "plan__product",
        "build",
        "manager",
    ).order_by("id")

    if ids:
        id_list = [int(x) for x in ids.split(",") if x.strip().isdigit()]
        if id_list:
            return queryset.filter(pk__in=id_list)

    if product_id:
        queryset = queryset.filter(plan__product_id=product_id)
    if plan_id:
        queryset = queryset.filter(plan_id=plan_id)
    if build_id:
        queryset = queryset.filter(build_id=build_id)

    return queryset


def _run_row(run):
    """Return a list of display values for a single test run."""
    total = getattr(run, "exec_count", 0)
    passed = getattr(run, "passed_count", 0)
    failed = getattr(run, "failed_count", 0)
    return [
        run.pk,
        run.summary,
        run.plan.name if run.plan else "",
        run.plan.product.name if run.plan and run.plan.product else "",
        run.build.name if run.build else "",
        run.manager.username if run.manager else "",
        run.start_date.strftime("%Y-%m-%d") if run.start_date else "",
        run.stop_date.strftime("%Y-%m-%d") if run.stop_date else "",
        total,
        passed,
        failed,
    ]


@method_decorator(login_required, name="dispatch")
class TestRunBrowserView(TemplateView):
    """
    Test Run Browser with tree navigation.
    Left panel: Product → TestPlan → TestRun tree
    Right panel: Selected test run details
    """

    template_name = "tcms_test_browser/run_browser.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        context["active_nav"] = "runs"

        runs_qs = TestRun.objects.select_related("build", "manager").order_by("summary")
        plans_qs = TestPlan.objects.prefetch_related(
            Prefetch("run", queryset=runs_qs)
        ).order_by("name")

        products = Product.objects.prefetch_related(
            Prefetch("plan", queryset=plans_qs)
        ).order_by("name")

        tree_data = []
        for product in products:
            plans = product.plan.all()
            plan_list = []
            product_run_count = 0
            for plan in plans:
                runs = plan.run.all()
                run_list = [
                    {
                        "id": run.pk,
                        "summary": run.summary,
                        "stop_date": run.stop_date.isoformat() if run.stop_date else None,
                    }
                    for run in runs
                ]
                product_run_count += len(run_list)
                plan_list.append(
                    {
                        "id": plan.pk,
                        "name": plan.name,
                        "count": len(run_list),
                        "runs": run_list,
                    }
                )

            tree_data.append(
                {
                    "id": product.pk,
                    "name": product.name,
                    "count": product_run_count,
                    "plans": plan_list,
                }
            )

        context["tree_data"] = tree_data
        return context


@login_required
def api_run_detail(request, run_id):
    """API endpoint to get test run details with executions."""
    try:
        run = TestRun.objects.select_related(
            "plan", "plan__product", "build", "manager", "default_tester",
        ).get(pk=run_id)
    except TestRun.DoesNotExist:
        return JsonResponse({"error": "Test run not found"}, status=404)

    executions = list(
        TestExecution.objects.filter(run=run)
        .select_related("case", "status", "tested_by")
        .order_by("case__summary")
        .values(
            "id",
            "case__id",
            "case__summary",
            "status__name",
            "status__color",
            "tested_by__username",
            "stop_date",
        )
    )

    data = {
        "id": run.pk,
        "summary": run.summary,
        "notes": run.notes or "",
        "plan_id": run.plan_id,
        "plan_name": run.plan.name if run.plan else None,
        "product": run.plan.product.name if run.plan and run.plan.product else None,
        "build": run.build.name if run.build else None,
        "manager": run.manager.username if run.manager else None,
        "default_tester": run.default_tester.username if run.default_tester else None,
        "start_date": run.start_date.isoformat() if run.start_date else None,
        "stop_date": run.stop_date.isoformat() if run.stop_date else None,
        "planned_start": run.planned_start.isoformat() if hasattr(run, "planned_start") and run.planned_start else None,
        "planned_stop": run.planned_stop.isoformat() if hasattr(run, "planned_stop") and run.planned_stop else None,
        "executions": [
            {
                "id": e["id"],
                "case_id": e["case__id"],
                "case_summary": e["case__summary"],
                "status": e["status__name"],
                "status_color": e["status__color"] or "",
                "tested_by": e["tested_by__username"],
                "stop_date": e["stop_date"].isoformat() if e["stop_date"] else None,
            }
            for e in executions
        ],
    }

    return JsonResponse(data)


@login_required
def api_search_runs(request):
    """API endpoint to search test runs."""
    query = request.GET.get("q", "").strip()
    product_id = request.GET.get("product")
    plan_id = request.GET.get("plan")
    build_id = request.GET.get("build")

    runs = TestRun.objects.select_related(
        "plan", "plan__product", "build", "manager",
    )

    if query:
        runs = runs.filter(summary__icontains=query)
    if product_id:
        runs = runs.filter(plan__product_id=product_id)
    if plan_id:
        runs = runs.filter(plan_id=plan_id)
    if build_id:
        runs = runs.filter(build_id=build_id)

    runs = runs.order_by("summary")[:100]

    data = [
        {
            "id": r.pk,
            "summary": r.summary,
            "plan": r.plan.name if r.plan else None,
            "product": r.plan.product.name if r.plan and r.plan.product else None,
            "build": r.build.name if r.build else None,
            "manager": r.manager.username if r.manager else None,
            "stop_date": r.stop_date.isoformat() if r.stop_date else None,
        }
        for r in runs
    ]

    return JsonResponse({"runs": data})


@login_required
def api_run_statistics(request):
    """API endpoint returning test run statistics for charts."""
    product_id = request.GET.get("product")

    run_qs = TestRun.objects.all()
    if product_id:
        run_qs = run_qs.filter(plan__product_id=product_id)

    total = run_qs.count()

    # Execution status distribution (aggregated from TestExecution)
    exec_qs = TestExecution.objects.all()
    if product_id:
        exec_qs = exec_qs.filter(run__plan__product_id=product_id)

    by_exec_status = list(
        exec_qs.values("status__name")
        .annotate(count=Count("id"))
        .order_by("status__name")
    )

    # Completed vs in-progress (based on stop_date presence)
    completed_count = run_qs.exclude(stop_date=None).count()
    in_progress_count = total - completed_count

    by_product = list(
        run_qs.values("plan__product__name")
        .annotate(count=Count("id"))
        .order_by("plan__product__name")
    )

    return JsonResponse(
        {
            "total": total,
            "by_exec_status": [
                {"name": item["status__name"] or "Unknown", "count": item["count"]}
                for item in by_exec_status
            ],
            "completion": {"completed": completed_count, "in_progress": in_progress_count},
            "by_product": [
                {"name": item["plan__product__name"] or "Unknown", "count": item["count"]}
                for item in by_product
            ],
        }
    )


@login_required
def api_browse_runs(request):
    """Paginated browse endpoint for test runs (chart drill-down)."""
    product_id = request.GET.get("product")
    page = int(request.GET.get("page", 1))
    page_size = int(request.GET.get("page_size", 25))

    queryset = TestRun.objects.select_related("plan", "plan__product").order_by("-id")

    if product_id:
        queryset = queryset.filter(plan__product_id=product_id)

    exec_status = request.GET.get("exec_status")
    if exec_status:
        run_ids = (
            TestExecution.objects.filter(status__name=exec_status)
            .values_list("run_id", flat=True)
            .distinct()
        )
        queryset = queryset.filter(pk__in=run_ids)

    product_name = request.GET.get("product_name")
    if product_name:
        queryset = queryset.filter(plan__product__name=product_name)

    total = queryset.count()
    total_pages = math.ceil(total / page_size) if total > 0 else 1
    page = max(1, min(page, total_pages))
    offset = (page - 1) * page_size

    runs = queryset[offset:offset + page_size]

    data = [
        {
            "id": r.pk,
            "summary": r.summary,
            "product": r.plan.product.name if r.plan and r.plan.product else None,
            "plan": r.plan.name if r.plan else None,
            "status": "Completed" if r.stop_date else "In Progress",
        }
        for r in runs
    ]

    return JsonResponse({
        "runs": data,
        "page": page,
        "total_pages": total_pages,
        "total": total,
    })


@login_required
def api_browse_executions(request):
    """Paginated browse endpoint for test executions (chart drill-down)."""
    product_id = request.GET.get("product")
    page = int(request.GET.get("page", 1))
    page_size = int(request.GET.get("page_size", 25))

    queryset = TestExecution.objects.select_related(
        "run", "case", "status",
    ).order_by("-pk")

    if product_id:
        queryset = queryset.filter(run__plan__product_id=product_id)

    exec_status = request.GET.get("exec_status")
    if exec_status:
        queryset = queryset.filter(status__name=exec_status)

    total = queryset.count()
    total_pages = math.ceil(total / page_size) if total > 0 else 1
    page = max(1, min(page, total_pages))
    offset = (page - 1) * page_size

    executions = queryset[offset:offset + page_size]

    data = [
        {
            "id": e.pk,
            "case_summary": e.case.summary if e.case else None,
            "run_summary": e.run.summary if e.run else None,
            "status": e.status.name if e.status else None,
        }
        for e in executions
    ]

    return JsonResponse({
        "executions": data,
        "page": page,
        "total_pages": total_pages,
        "total": total,
    })


@login_required
def api_run_report(request):
    """API endpoint to download a CSV report of test runs."""
    queryset = _get_run_report_queryset(request)
    products, plans_list = _extract_run_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Runs", products, plans_list)

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = 'attachment; filename="test_runs_report.csv"'

    writer = csv.writer(response)
    _write_csv_metadata(writer, metadata)
    writer.writerow(RUN_REPORT_HEADERS)

    for run in queryset:
        execs = TestExecution.objects.filter(run=run)
        total_exec = execs.count()
        passed = execs.filter(status__weight__gt=0).count()
        failed = execs.filter(status__weight__lt=0).count()
        writer.writerow([
            run.pk,
            run.summary,
            run.plan.name if run.plan else "",
            run.plan.product.name if run.plan and run.plan.product else "",
            run.build.name if run.build else "",
            run.manager.username if run.manager else "",
            run.start_date.strftime("%Y-%m-%d") if run.start_date else "",
            run.stop_date.strftime("%Y-%m-%d") if run.stop_date else "",
            total_exec,
            passed,
            failed,
        ])

    return response


@login_required
def api_run_report_excel(request):
    """API endpoint to download an Excel report of test runs."""
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter

    queryset = _get_run_report_queryset(request)
    products, plans_list = _extract_run_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Runs", products, plans_list)

    wb = Workbook()

    # Create Overview sheet with charts (first sheet)
    chart_data = _compute_run_chart_data(queryset)
    if any(chart_data.values()):
        ws_charts = wb.active
        ws_charts.title = "Overview"

        # Add charts horizontally
        if chart_data["by_exec_status"]:
            _add_excel_pie_chart(ws_charts, "Execution Status", chart_data["by_exec_status"], "A1", data_col=20)
        if chart_data["completion"]:
            _add_excel_pie_chart(ws_charts, "Completion Status", chart_data["completion"], "J1", data_col=22)

        # Create data sheet
        ws = wb.create_sheet("Test Runs")
    else:
        ws = wb.active
        ws.title = "Test Runs"

    # Metadata rows
    data_start = _write_excel_metadata(ws, metadata)

    # Build data rows
    rows = []
    for run in queryset:
        execs = TestExecution.objects.filter(run=run)
        total_exec = execs.count()
        passed = execs.filter(status__weight__gt=0).count()
        failed = execs.filter(status__weight__lt=0).count()
        rows.append([
            run.pk,
            run.summary,
            run.plan.name if run.plan else "",
            run.plan.product.name if run.plan and run.plan.product else "",
            run.build.name if run.build else "",
            run.manager.username if run.manager else "",
            run.start_date.strftime("%Y-%m-%d") if run.start_date else "",
            run.stop_date.strftime("%Y-%m-%d") if run.stop_date else "",
            total_exec,
            passed,
            failed,
        ])

    # Write styled table
    last_row = _write_excel_table(ws, RUN_REPORT_HEADERS, rows, start_row=data_start)

    # Auto-filter on all columns
    if last_row > data_start:
        ws.auto_filter.ref = (
            f"A{data_start}:{get_column_letter(len(RUN_REPORT_HEADERS))}{last_row - 1}"
        )

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = 'attachment; filename="test_runs_report.xlsx"'
    return response


@login_required
def api_run_report_docx(request):
    """API endpoint to download a Word document report of test runs."""
    from docx import Document
    from docx.shared import Inches

    queryset = _get_run_report_queryset(request)
    products, plans_list = _extract_run_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Runs", products, plans_list)

    doc = Document()
    doc.add_heading("Test Runs Report", level=1)
    _write_docx_metadata(doc, metadata)

    # Add charts if available
    chart_data = _compute_run_chart_data(queryset)
    chart_items = [
        (t, d)
        for t, d in [
            ("Execution Status", chart_data.get("by_exec_status")),
            ("Completion Status", chart_data.get("completion")),
        ]
        if d
    ]
    if chart_items:
        doc.add_heading("Statistics Overview", level=2)
        _add_docx_charts_row(doc, chart_items)
        doc.add_paragraph()

    # Build data rows
    rows = []
    for tr in queryset:
        execs = TestExecution.objects.filter(run=tr)
        total_exec = execs.count()
        passed = execs.filter(status__weight__gt=0).count()
        failed = execs.filter(status__weight__lt=0).count()
        rows.append([
            tr.pk,
            tr.summary,
            tr.plan.name if tr.plan else "",
            tr.plan.product.name if tr.plan and tr.plan.product else "",
            tr.build.name if tr.build else "",
            tr.manager.username if tr.manager else "",
            tr.start_date.strftime("%Y-%m-%d") if tr.start_date else "",
            tr.stop_date.strftime("%Y-%m-%d") if tr.stop_date else "",
            total_exec,
            passed,
            failed,
        ])

    # Write styled table
    _write_docx_table(doc, None, RUN_REPORT_HEADERS, rows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    response["Content-Disposition"] = 'attachment; filename="test_runs_report.docx"'
    return response


@login_required
def api_run_report_pdf(request):
    """API endpoint to download a PDF report of test runs."""
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table

    queryset = _get_run_report_queryset(request)
    products, plans_list = _extract_run_metadata_from_qs(queryset)
    metadata = _build_report_metadata(request, "Test Runs", products, plans_list)
    styles = getSampleStyleSheet()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(letter), topMargin=0.5 * inch)
    elements = []

    elements.append(Paragraph("Test Runs Report", styles["Title"]))
    elements.append(Spacer(1, 12))
    elements.extend(_build_pdf_metadata_elements(metadata, styles))

    # Add charts if available
    chart_data = _compute_run_chart_data(queryset)
    chart_items = [
        (t, d)
        for t, d in [
            ("Execution Status", chart_data.get("by_exec_status")),
            ("Completion Status", chart_data.get("completion")),
        ]
        if d
    ]
    if chart_items:
        elements.append(Paragraph("Statistics Overview", styles["Heading2"]))
        elements.append(Spacer(1, 12))
        charts_table = _add_pdf_charts_row(chart_items)
        if charts_table:
            elements.append(charts_table)
        elements.append(Spacer(1, 12))

    cell_style = styles["Normal"]
    cell_style.fontSize = 7
    cell_style.leading = 9

    table_data = [RUN_REPORT_HEADERS]
    for tr in queryset:
        execs = TestExecution.objects.filter(run=tr)
        total_exec = execs.count()
        passed = execs.filter(status__weight__gt=0).count()
        failed = execs.filter(status__weight__lt=0).count()
        row = [
            tr.pk,
            Paragraph(str(tr.summary), cell_style),
            tr.plan.name if tr.plan else "",
            tr.plan.product.name if tr.plan and tr.plan.product else "",
            tr.build.name if tr.build else "",
            tr.manager.username if tr.manager else "",
            tr.start_date.strftime("%Y-%m-%d") if tr.start_date else "",
            tr.stop_date.strftime("%Y-%m-%d") if tr.stop_date else "",
            total_exec,
            passed,
            failed,
        ]
        table_data.append(row)

    col_widths = [
        0.4 * inch,   # ID
        2.0 * inch,   # Summary
        1.2 * inch,   # Plan
        0.9 * inch,   # Product
        0.7 * inch,   # Build
        0.7 * inch,   # Manager
        0.7 * inch,   # Started
        0.7 * inch,   # Stopped
        0.6 * inch,   # Total
        0.5 * inch,   # Passed
        0.5 * inch,   # Failed
    ]

    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    table.setStyle(_get_pdf_table_style())
    elements.append(table)

    doc.build(elements)
    buf.seek(0)

    response = HttpResponse(buf.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="test_runs_report.pdf"'
    return response


# ==========================================
# Consolidated Browser
# ==========================================


@method_decorator(login_required, name="dispatch")
class ConsolidatedBrowserView(TemplateView):
    """
    Consolidated Browser with Dashboard, Plan-centric, and Case-centric tabs.
    """

    template_name = "tcms_test_browser/consolidated.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["active_nav"] = "consolidated"
        context["products"] = Product.objects.order_by("name")
        return context


@login_required
def api_consolidated_dashboard(request):
    """Dashboard data: totals, coverage gaps, recent activity."""
    product_id = request.GET.get("product")
    runs_page = int(request.GET.get("runs_page", 1))
    cases_page = int(request.GET.get("cases_page", 1))
    cases_no_plan_page = int(request.GET.get("cases_no_plan_page", 1))
    plans_no_run_page = int(request.GET.get("plans_no_run_page", 1))
    page_size = int(request.GET.get("page_size", 10))

    case_qs = TestCase.objects.all()
    plan_qs = TestPlan.objects.all()
    run_qs = TestRun.objects.all()
    exec_qs = TestExecution.objects.all()

    if product_id:
        case_qs = case_qs.filter(category__product_id=product_id)
        plan_qs = plan_qs.filter(product_id=product_id)
        run_qs = run_qs.filter(plan__product_id=product_id)
        exec_qs = exec_qs.filter(run__plan__product_id=product_id)

    total_cases = case_qs.count()
    total_plans = plan_qs.count()
    total_runs = run_qs.count()

    # Execution pass/fail rates
    total_execs = exec_qs.count()
    passed_execs = exec_qs.filter(status__weight__gt=0).count()
    failed_execs = exec_qs.filter(status__weight__lt=0).count()

    # Coverage gaps: cases not in any plan (with pagination)
    cases_no_plan_qs = case_qs.filter(plan=None)
    cases_no_plan_total = cases_no_plan_qs.count()
    cases_no_plan_total_pages = math.ceil(cases_no_plan_total / page_size) if cases_no_plan_total > 0 else 1
    cases_no_plan_page = max(1, min(cases_no_plan_page, cases_no_plan_total_pages))
    cases_no_plan_offset = (cases_no_plan_page - 1) * page_size

    cases_without_plans = (
        cases_no_plan_qs
        .values("id", "summary")
        .order_by("-id")[cases_no_plan_offset:cases_no_plan_offset + page_size]
    )

    # Plans without runs (with pagination)
    plans_no_run_qs = plan_qs.filter(run=None)
    plans_no_run_total = plans_no_run_qs.count()
    plans_no_run_total_pages = math.ceil(plans_no_run_total / page_size) if plans_no_run_total > 0 else 1
    plans_no_run_page = max(1, min(plans_no_run_page, plans_no_run_total_pages))
    plans_no_run_offset = (plans_no_run_page - 1) * page_size

    plans_without_runs = (
        plans_no_run_qs
        .values("id", "name")
        .order_by("-id")[plans_no_run_offset:plans_no_run_offset + page_size]
    )

    # Recent activity with pagination
    runs_total_pages = math.ceil(total_runs / page_size) if total_runs > 0 else 1
    runs_page = max(1, min(runs_page, runs_total_pages))
    runs_offset = (runs_page - 1) * page_size

    recent_runs = list(
        run_qs.select_related("plan", "plan__product", "manager")
        .order_by("-id")[runs_offset:runs_offset + page_size]
        .values(
            "id", "summary", "plan__name", "plan__product__name",
            "manager__username", "start_date", "stop_date",
        )
    )

    cases_total_pages = math.ceil(total_cases / page_size) if total_cases > 0 else 1
    cases_page = max(1, min(cases_page, cases_total_pages))
    cases_offset = (cases_page - 1) * page_size

    recent_cases = list(
        case_qs.select_related("category", "category__product", "author")
        .order_by("-id")[cases_offset:cases_offset + page_size]
        .values(
            "id", "summary", "category__product__name",
            "author__username", "create_date",
        )
    )

    # Execution status breakdown
    by_exec_status = list(
        exec_qs.values("status__name")
        .annotate(count=Count("id"))
        .order_by("status__name")
    )

    return JsonResponse({
        "totals": {
            "cases": total_cases,
            "plans": total_plans,
            "runs": total_runs,
            "executions": total_execs,
        },
        "execution_rates": {
            "passed": passed_execs,
            "failed": failed_execs,
            "other": total_execs - passed_execs - failed_execs,
            "total": total_execs,
        },
        "by_exec_status": [
            {"name": item["status__name"] or "Unknown", "count": item["count"]}
            for item in by_exec_status
        ],
        "cases_without_plans": list(cases_without_plans),
        "cases_without_plans_pagination": {
            "page": cases_no_plan_page,
            "total_pages": cases_no_plan_total_pages,
            "total": cases_no_plan_total,
        },
        "plans_without_runs": list(plans_without_runs),
        "plans_without_runs_pagination": {
            "page": plans_no_run_page,
            "total_pages": plans_no_run_total_pages,
            "total": plans_no_run_total,
        },
        "recent_runs": [
            {
                "id": r["id"],
                "summary": r["summary"],
                "plan": r["plan__name"],
                "product": r["plan__product__name"],
                "manager": r["manager__username"],
                "start_date": r["start_date"].isoformat() if r["start_date"] else None,
                "stop_date": r["stop_date"].isoformat() if r["stop_date"] else None,
            }
            for r in recent_runs
        ],
        "recent_runs_pagination": {
            "page": runs_page,
            "total_pages": runs_total_pages,
            "total": total_runs,
        },
        "recent_cases": [
            {
                "id": r["id"],
                "summary": r["summary"],
                "product": r["category__product__name"],
                "author": r["author__username"],
                "create_date": r["create_date"].isoformat() if r["create_date"] else None,
            }
            for r in recent_cases
        ],
        "recent_cases_pagination": {
            "page": cases_page,
            "total_pages": cases_total_pages,
            "total": total_cases,
        },
    })


@login_required
def api_consolidated_sankey(request):
    """Return Sankey diagram data: TestCase -> TestPlan -> TestRun links."""
    product_id = request.GET.get("product")

    plan_qs = TestPlan.objects.all()
    run_qs = TestRun.objects.all()

    if product_id:
        plan_qs = plan_qs.filter(product_id=product_id)
        run_qs = run_qs.filter(plan__product_id=product_id)

    # Build nodes and links
    # Nodes: products (left), plans (middle), runs (right)
    nodes = []
    node_index = {}
    links = []

    # Get plans with their case counts and product info
    plans = list(
        plan_qs.select_related("product")
        .annotate(case_count=Count("cases", distinct=True))
        .values("id", "name", "product__name", "product__id", "case_count")
        .order_by("product__name", "name")
    )

    # Get runs with their plan info
    runs = list(
        run_qs.select_related("plan", "plan__product")
        .values("id", "summary", "plan__id", "plan__name")
        .order_by("plan__id", "-id")
    )

    # Build product nodes (source)
    products_seen = {}
    for p in plans:
        prod_name = p["product__name"] or "Unknown"
        if prod_name not in products_seen:
            idx = len(nodes)
            nodes.append({"name": prod_name, "type": "product"})
            node_index["product:" + prod_name] = idx
            products_seen[prod_name] = idx

    # Build plan nodes (middle)
    for p in plans:
        plan_key = "plan:" + str(p["id"])
        idx = len(nodes)
        nodes.append({"name": "TP-" + str(p["id"]) + ": " + (p["name"] or ""), "type": "plan"})
        node_index[plan_key] = idx

        # Link product -> plan (weight = case_count or 1 if 0)
        prod_key = "product:" + (p["product__name"] or "Unknown")
        if prod_key in node_index:
            links.append({
                "source": node_index[prod_key],
                "target": idx,
                "value": max(p["case_count"], 1),
            })

    # Build run nodes (right) and plan -> run links
    for r in runs:
        run_key = "run:" + str(r["id"])
        idx = len(nodes)
        nodes.append({"name": "TR-" + str(r["id"]) + ": " + (r["summary"] or ""), "type": "run"})
        node_index[run_key] = idx

        plan_key = "plan:" + str(r["plan__id"])
        if plan_key in node_index:
            # Weight = 1 per run
            links.append({
                "source": node_index[plan_key],
                "target": idx,
                "value": 1,
            })

    return JsonResponse({
        "nodes": nodes,
        "links": links,
    })


@login_required
def api_consolidated_plan_detail(request, plan_id):
    """Plan-centric drill-down: cases with execution status across runs."""
    try:
        plan = TestPlan.objects.select_related(
            "product", "product_version", "type", "author",
        ).get(pk=plan_id)
    except TestPlan.DoesNotExist:
        return JsonResponse({"error": "Test plan not found"}, status=404)

    # All cases in this plan
    cases = list(
        plan.cases.select_related("case_status", "priority")
        .order_by("id")
        .values("id", "summary", "case_status__name", "priority__value")
    )

    # All runs for this plan
    runs = list(
        plan.run.order_by("-id")
        .values("id", "summary", "start_date", "stop_date")
    )
    run_ids = [r["id"] for r in runs]
    case_ids = [c["id"] for c in cases]

    # All executions for these cases in these runs
    executions = list(
        TestExecution.objects.filter(run_id__in=run_ids, case_id__in=case_ids)
        .select_related("status")
        .values("case_id", "run_id", "status__name", "status__color", "status__weight")
    )

    # Build execution map: case_id -> run_id -> status
    exec_map = {}
    for e in executions:
        case_id = e["case_id"]
        run_id = e["run_id"]
        if case_id not in exec_map:
            exec_map[case_id] = {}
        exec_map[case_id][run_id] = {
            "status": e["status__name"],
            "color": e["status__color"] or "",
            "weight": e["status__weight"],
        }

    # Per-case aggregated stats
    cases_with_stats = []
    for c in cases:
        cid = c["id"]
        case_execs = exec_map.get(cid, {})
        total = len(case_execs)
        passed = sum(1 for e in case_execs.values() if e["weight"] and e["weight"] > 0)
        failed = sum(1 for e in case_execs.values() if e["weight"] and e["weight"] < 0)
        cases_with_stats.append({
            "id": cid,
            "summary": c["summary"],
            "status": c["case_status__name"],
            "priority": c["priority__value"],
            "total_executions": total,
            "passed": passed,
            "failed": failed,
            "exec_by_run": {
                str(rid): case_execs.get(rid, None)
                for rid in run_ids
            },
        })

    # Per-run completion percentages
    runs_with_stats = []
    for r in runs:
        rid = r["id"]
        run_execs = [
            exec_map.get(cid, {}).get(rid)
            for cid in case_ids
        ]
        total = sum(1 for e in run_execs if e is not None)
        passed = sum(1 for e in run_execs if e and e["weight"] and e["weight"] > 0)
        runs_with_stats.append({
            "id": rid,
            "summary": r["summary"],
            "start_date": r["start_date"].isoformat() if r["start_date"] else None,
            "stop_date": r["stop_date"].isoformat() if r["stop_date"] else None,
            "total_executions": total,
            "passed": passed,
            "completion_pct": round(passed / total * 100) if total > 0 else 0,
        })

    return JsonResponse({
        "plan": {
            "id": plan.pk,
            "name": plan.name,
            "product": plan.product.name if plan.product else None,
            "version": plan.product_version.value if plan.product_version else None,
            "type": plan.type.name if plan.type else None,
            "author": plan.author.username if plan.author else None,
            "is_active": plan.is_active,
        },
        "cases": cases_with_stats,
        "runs": runs_with_stats,
    })


@login_required
def api_consolidated_case_detail(request, case_id):
    """Case-centric drill-down: all plans and executions across runs."""
    try:
        tc = TestCase.objects.select_related(
            "case_status", "priority", "author", "category", "category__product",
        ).get(pk=case_id)
    except TestCase.DoesNotExist:
        return JsonResponse({"error": "Test case not found"}, status=404)

    # All plans containing this case
    plans = list(
        tc.plan.select_related("product")
        .order_by("id")
        .values("id", "name", "product__name", "is_active")
    )

    # All executions of this case across all runs
    executions = list(
        TestExecution.objects.filter(case=tc)
        .select_related("run", "run__plan", "status", "tested_by")
        .order_by("-run__id")
        .values(
            "id", "run_id", "run__summary", "run__plan__name",
            "status__name", "status__color", "status__weight",
            "tested_by__username", "stop_date",
        )
    )

    # Group executions by run
    runs_map = {}
    for e in executions:
        rid = e["run_id"]
        if rid not in runs_map:
            runs_map[rid] = {
                "run_id": rid,
                "run_summary": e["run__summary"],
                "plan_name": e["run__plan__name"],
                "executions": [],
            }
        runs_map[rid]["executions"].append({
            "id": e["id"],
            "status": e["status__name"],
            "status_color": e["status__color"] or "",
            "weight": e["status__weight"],
            "tested_by": e["tested_by__username"],
            "stop_date": e["stop_date"].isoformat() if e["stop_date"] else None,
        })

    # Execution timeline (flat list, most recent first)
    timeline = [
        {
            "id": e["id"],
            "run_id": e["run_id"],
            "run_summary": e["run__summary"],
            "status": e["status__name"],
            "status_color": e["status__color"] or "",
            "tested_by": e["tested_by__username"],
            "stop_date": e["stop_date"].isoformat() if e["stop_date"] else None,
        }
        for e in executions
    ]

    return JsonResponse({
        "case": {
            "id": tc.pk,
            "summary": tc.summary,
            "status": tc.case_status.name if tc.case_status else None,
            "priority": tc.priority.value if tc.priority else None,
            "author": tc.author.username if tc.author else None,
            "product": tc.category.product.name if tc.category and tc.category.product else None,
            "category": tc.category.name if tc.category else None,
        },
        "plans": list(plans),
        "executions_by_run": list(runs_map.values()),
        "timeline": timeline,
    })


# ==========================================
# Consolidated Dashboard Exports
# ==========================================

DASH_EXEC_STATUS_HEADERS = ["Status", "Count"]
DASH_CASES_NO_PLAN_HEADERS = ["ID", "Summary"]
DASH_PLANS_NO_RUN_HEADERS = ["ID", "Name"]
DASH_RECENT_RUNS_HEADERS = ["ID", "Summary", "Plan", "Product", "Manager", "Started", "Stopped"]
DASH_RECENT_CASES_HEADERS = ["ID", "Summary", "Product", "Author", "Created"]


def _get_dashboard_data(request):
    """Gather all dashboard data for export, mirroring api_consolidated_dashboard."""
    product_id = request.GET.get("product")

    case_qs = TestCase.objects.all()
    plan_qs = TestPlan.objects.all()
    run_qs = TestRun.objects.all()
    exec_qs = TestExecution.objects.all()

    if product_id:
        case_qs = case_qs.filter(category__product_id=product_id)
        plan_qs = plan_qs.filter(product_id=product_id)
        run_qs = run_qs.filter(plan__product_id=product_id)
        exec_qs = exec_qs.filter(run__plan__product_id=product_id)

    total_cases = case_qs.count()
    total_plans = plan_qs.count()
    total_runs = run_qs.count()
    total_execs = exec_qs.count()
    passed_execs = exec_qs.filter(status__weight__gt=0).count()
    failed_execs = exec_qs.filter(status__weight__lt=0).count()
    pass_rate = round(passed_execs / total_execs * 100) if total_execs > 0 else 0

    summary = {
        "cases": total_cases,
        "plans": total_plans,
        "runs": total_runs,
        "executions": total_execs,
        "passed": passed_execs,
        "failed": failed_execs,
        "pass_rate": pass_rate,
    }

    by_exec_status = list(
        exec_qs.values("status__name")
        .annotate(count=Count("id"))
        .order_by("status__name")
    )
    exec_status_rows = [
        [item["status__name"] or "Unknown", item["count"]]
        for item in by_exec_status
    ]

    cases_no_plan = list(
        case_qs.filter(plan=None)
        .values("id", "summary")
        .order_by("-id")[:20]
    )
    cases_no_plan_rows = [[c["id"], c["summary"]] for c in cases_no_plan]

    plans_no_run = list(
        plan_qs.filter(run=None)
        .values("id", "name")
        .order_by("-id")[:20]
    )
    plans_no_run_rows = [[p["id"], p["name"]] for p in plans_no_run]

    recent_runs = list(
        run_qs.select_related("plan", "plan__product", "manager")
        .order_by("-id")[:10]
        .values(
            "id", "summary", "plan__name", "plan__product__name",
            "manager__username", "start_date", "stop_date",
        )
    )
    recent_runs_rows = [
        [
            r["id"],
            r["summary"],
            r["plan__name"] or "",
            r["plan__product__name"] or "",
            r["manager__username"] or "",
            r["start_date"].strftime("%Y-%m-%d") if r["start_date"] else "",
            r["stop_date"].strftime("%Y-%m-%d") if r["stop_date"] else "",
        ]
        for r in recent_runs
    ]

    recent_cases = list(
        case_qs.select_related("category", "category__product", "author")
        .order_by("-id")[:10]
        .values(
            "id", "summary", "category__product__name",
            "author__username", "create_date",
        )
    )
    recent_cases_rows = [
        [
            c["id"],
            c["summary"],
            c["category__product__name"] or "",
            c["author__username"] or "",
            c["create_date"].strftime("%Y-%m-%d") if c["create_date"] else "",
        ]
        for c in recent_cases
    ]

    # Product/plan metadata for report headers
    products = list(
        plan_qs.values_list("product__name", flat=True)
        .distinct()
        .order_by("product__name")
    )
    products = [p for p in products if p]

    plans = list(
        plan_qs.values_list("pk", "name")
        .distinct()
        .order_by("pk")
    )
    plans_meta = ["TP-{}: {}".format(pk, name) for pk, name in plans]

    return (
        summary, exec_status_rows, cases_no_plan_rows,
        plans_no_run_rows, recent_runs_rows, recent_cases_rows,
        products, plans_meta,
    )


@login_required
def api_consolidated_dashboard_report(request):
    """CSV export for consolidated dashboard."""
    (
        summary, exec_status_rows, cases_no_plan_rows,
        plans_no_run_rows, recent_runs_rows, recent_cases_rows,
        products, plans_meta,
    ) = _get_dashboard_data(request)
    metadata = _build_report_metadata(
        request, "Consolidated Dashboard", products, plans_meta
    )

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = (
        'attachment; filename="consolidated_dashboard.csv"'
    )

    writer = csv.writer(response)
    _write_csv_metadata(writer, metadata)
    writer.writerow(["Consolidated Dashboard Report"])
    writer.writerow([])
    writer.writerow(["Summary"])
    writer.writerow(["Test Cases", summary["cases"]])
    writer.writerow(["Test Plans", summary["plans"]])
    writer.writerow(["Test Runs", summary["runs"]])
    writer.writerow(["Total Executions", summary["executions"]])
    writer.writerow(["Passed", summary["passed"]])
    writer.writerow(["Failed", summary["failed"]])
    writer.writerow(["Pass Rate", "{}%".format(summary["pass_rate"])])
    writer.writerow([])
    writer.writerow(["Execution Status Breakdown"])
    writer.writerow(DASH_EXEC_STATUS_HEADERS)
    for row in exec_status_rows:
        writer.writerow(row)
    writer.writerow([])
    writer.writerow(["Cases Without Plans"])
    writer.writerow(DASH_CASES_NO_PLAN_HEADERS)
    for row in cases_no_plan_rows:
        writer.writerow(row)
    writer.writerow([])
    writer.writerow(["Plans Without Runs"])
    writer.writerow(DASH_PLANS_NO_RUN_HEADERS)
    for row in plans_no_run_rows:
        writer.writerow(row)
    writer.writerow([])
    writer.writerow(["Recent Test Runs"])
    writer.writerow(DASH_RECENT_RUNS_HEADERS)
    for row in recent_runs_rows:
        writer.writerow(row)
    writer.writerow([])
    writer.writerow(["Recent Test Cases"])
    writer.writerow(DASH_RECENT_CASES_HEADERS)
    for row in recent_cases_rows:
        writer.writerow(row)

    return response


@login_required
def api_consolidated_dashboard_report_excel(request):
    """Excel export for consolidated dashboard."""
    from openpyxl import Workbook

    (
        summary, exec_status_rows, cases_no_plan_rows,
        plans_no_run_rows, recent_runs_rows, recent_cases_rows,
        products, plans_meta,
    ) = _get_dashboard_data(request)
    metadata = _build_report_metadata(
        request, "Consolidated Dashboard", products, plans_meta
    )

    wb = Workbook()

    # Report Info sheet
    ws_info = wb.active
    ws_info.title = "Report Info"
    _write_excel_metadata(ws_info, metadata)
    for col in ws_info.columns:
        max_length = 0
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws_info.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)

    # Summary sheet
    ws_summary = wb.create_sheet("Summary")
    summary_headers = ["Metric", "Value"]
    summary_rows = [
        ["Test Cases", summary["cases"]],
        ["Test Plans", summary["plans"]],
        ["Test Runs", summary["runs"]],
        ["Total Executions", summary["executions"]],
        ["Passed", summary["passed"]],
        ["Failed", summary["failed"]],
        ["Pass Rate", "{}%".format(summary["pass_rate"])],
    ]
    _write_excel_table(ws_summary, summary_headers, summary_rows, start_row=1)

    # Add execution status chart to Summary sheet
    if exec_status_rows:
        exec_status_dict = {row[0]: row[1] for row in exec_status_rows}
        _add_excel_pie_chart(ws_summary, "Execution Status Breakdown", exec_status_dict, "E1", data_col=20)

    # Execution Status sheet
    ws_exec = wb.create_sheet("Execution Status")
    _write_excel_table(ws_exec, DASH_EXEC_STATUS_HEADERS, exec_status_rows, start_row=1)

    # Coverage Gaps sheet
    ws_gaps = wb.create_sheet("Coverage Gaps")
    gap_headers = ["Type", "ID", "Name/Summary"]
    gap_rows = []
    for row in cases_no_plan_rows:
        gap_rows.append(["Case Without Plan", row[0], row[1]])
    for row in plans_no_run_rows:
        gap_rows.append(["Plan Without Run", row[0], row[1]])
    _write_excel_table(ws_gaps, gap_headers, gap_rows, start_row=1)

    # Recent Runs sheet
    ws_runs = wb.create_sheet("Recent Runs")
    _write_excel_table(ws_runs, DASH_RECENT_RUNS_HEADERS, recent_runs_rows, start_row=1)

    # Recent Cases sheet
    ws_cases = wb.create_sheet("Recent Cases")
    _write_excel_table(ws_cases, DASH_RECENT_CASES_HEADERS, recent_cases_rows, start_row=1)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = (
        'attachment; filename="consolidated_dashboard.xlsx"'
    )
    return response


@login_required
def api_consolidated_dashboard_report_docx(request):
    """Word export for consolidated dashboard."""
    from docx import Document

    (
        summary, exec_status_rows, cases_no_plan_rows,
        plans_no_run_rows, recent_runs_rows, recent_cases_rows,
        products, plans_meta,
    ) = _get_dashboard_data(request)
    metadata = _build_report_metadata(
        request, "Consolidated Dashboard", products, plans_meta
    )

    doc = Document()
    doc.add_heading("Consolidated Dashboard Report", level=1)
    _write_docx_metadata(doc, metadata)

    # Add execution status chart if available
    if exec_status_rows:
        exec_status_dict = {row[0]: row[1] for row in exec_status_rows}
        doc.add_heading("Statistics Overview", level=2)
        _add_docx_charts_row(doc, [("Execution Status", exec_status_dict)])
        doc.add_paragraph()

    # Summary
    summary_rows = [
        ["Test Cases", summary["cases"]],
        ["Test Plans", summary["plans"]],
        ["Test Runs", summary["runs"]],
        ["Total Executions", summary["executions"]],
        ["Passed", summary["passed"]],
        ["Failed", summary["failed"]],
        ["Pass Rate", "{}%".format(summary["pass_rate"])],
    ]
    _write_docx_table(doc, "Summary", ["Metric", "Value"], summary_rows)
    _write_docx_table(doc, "Execution Status Breakdown", DASH_EXEC_STATUS_HEADERS, exec_status_rows)
    _write_docx_table(doc, "Cases Without Plans", DASH_CASES_NO_PLAN_HEADERS, cases_no_plan_rows)
    _write_docx_table(doc, "Plans Without Runs", DASH_PLANS_NO_RUN_HEADERS, plans_no_run_rows)
    _write_docx_table(doc, "Recent Test Runs", DASH_RECENT_RUNS_HEADERS, recent_runs_rows)
    _write_docx_table(doc, "Recent Test Cases", DASH_RECENT_CASES_HEADERS, recent_cases_rows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    response["Content-Disposition"] = (
        'attachment; filename="consolidated_dashboard.docx"'
    )
    return response


@login_required
def api_consolidated_dashboard_report_pdf(request):
    """PDF export for consolidated dashboard."""
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table

    (
        summary, exec_status_rows, cases_no_plan_rows,
        plans_no_run_rows, recent_runs_rows, recent_cases_rows,
        products, plans_meta,
    ) = _get_dashboard_data(request)
    metadata = _build_report_metadata(
        request, "Consolidated Dashboard", products, plans_meta
    )

    styles = getSampleStyleSheet()
    cell_style = styles["Normal"]
    cell_style.fontSize = 7
    cell_style.leading = 9

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(letter), topMargin=0.5 * inch)
    elements = []

    elements.append(Paragraph("Consolidated Dashboard Report", styles["Title"]))
    elements.append(Spacer(1, 12))
    elements.extend(_build_pdf_metadata_elements(metadata, styles))

    # Add execution status chart if available
    if exec_status_rows:
        exec_status_dict = {row[0]: row[1] for row in exec_status_rows}
        elements.append(Paragraph("Statistics Overview", styles["Heading2"]))
        elements.append(Spacer(1, 12))
        charts_table = _add_pdf_charts_row([("Execution Status", exec_status_dict)])
        if charts_table:
            elements.append(charts_table)
        elements.append(Spacer(1, 18))

    table_style = _get_pdf_table_style()

    def add_section(title, headers, rows, col_widths):
        elements.append(Paragraph(title, styles["Heading2"]))
        elements.append(Spacer(1, 6))
        data = [headers]
        for row in rows:
            r = list(row)
            # Wrap long text columns in Paragraph
            for i, val in enumerate(r):
                if isinstance(val, str) and len(val) > 30:
                    r[i] = Paragraph(val, cell_style)
            data.append(r)
        if len(data) > 1:
            t = Table(data, colWidths=col_widths, repeatRows=1)
            t.setStyle(table_style)
            elements.append(t)
        else:
            elements.append(Paragraph("<i>No data</i>", cell_style))
        elements.append(Spacer(1, 18))

    # Summary
    summary_rows = [
        ["Test Cases", summary["cases"]],
        ["Test Plans", summary["plans"]],
        ["Test Runs", summary["runs"]],
        ["Total Executions", summary["executions"]],
        ["Passed", summary["passed"]],
        ["Failed", summary["failed"]],
        ["Pass Rate", "{}%".format(summary["pass_rate"])],
    ]
    add_section("Summary", ["Metric", "Value"], summary_rows,
                [2.0 * inch, 1.5 * inch])

    add_section("Execution Status Breakdown", DASH_EXEC_STATUS_HEADERS, exec_status_rows,
                [2.0 * inch, 1.0 * inch])

    add_section("Cases Without Plans", DASH_CASES_NO_PLAN_HEADERS, cases_no_plan_rows,
                [0.6 * inch, 4.0 * inch])

    add_section("Plans Without Runs", DASH_PLANS_NO_RUN_HEADERS, plans_no_run_rows,
                [0.6 * inch, 4.0 * inch])

    add_section("Recent Test Runs", DASH_RECENT_RUNS_HEADERS, recent_runs_rows,
                [0.4 * inch, 2.0 * inch, 1.2 * inch, 1.0 * inch,
                 0.8 * inch, 0.8 * inch, 0.8 * inch])

    add_section("Recent Test Cases", DASH_RECENT_CASES_HEADERS, recent_cases_rows,
                [0.4 * inch, 2.5 * inch, 1.2 * inch, 0.8 * inch, 0.8 * inch])

    doc.build(elements)
    buf.seek(0)

    response = HttpResponse(buf.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = (
        'attachment; filename="consolidated_dashboard.pdf"'
    )
    return response


# ==========================================
# Consolidated Plan Drill-Down Exports
# ==========================================


def _get_plan_drilldown_data(plan_id):
    """Gather runs and cases data for a plan drill-down export."""
    plan = TestPlan.objects.select_related(
        "product", "product_version", "type", "author",
    ).get(pk=plan_id)

    cases = list(
        plan.cases.select_related("case_status")
        .order_by("id")
        .values("id", "summary", "case_status__name")
    )
    runs = list(
        plan.run.order_by("-id")
        .values("id", "summary", "start_date", "stop_date")
    )
    run_ids = [r["id"] for r in runs]
    case_ids = [c["id"] for c in cases]

    executions = list(
        TestExecution.objects.filter(run_id__in=run_ids, case_id__in=case_ids)
        .values("case_id", "run_id", "status__weight")
    )

    exec_map = {}
    for e in executions:
        exec_map.setdefault(e["case_id"], {}).setdefault(e["run_id"], []).append(
            e["status__weight"]
        )

    runs_rows = []
    for r in runs:
        rid = r["id"]
        run_execs = TestExecution.objects.filter(run_id=rid)
        total = run_execs.count()
        passed = run_execs.filter(status__weight__gt=0).count()
        pct = round(passed / total * 100) if total > 0 else 0
        runs_rows.append([
            rid,
            r["summary"],
            r["start_date"].strftime("%Y-%m-%d") if r["start_date"] else "",
            r["stop_date"].strftime("%Y-%m-%d") if r["stop_date"] else "",
            total,
            passed,
            "{}%".format(pct),
        ])

    cases_rows = []
    for c in cases:
        cid = c["id"]
        case_execs = exec_map.get(cid, {})
        total = sum(len(v) for v in case_execs.values())
        passed = sum(
            1 for weights in case_execs.values()
            for w in weights if w and w > 0
        )
        failed = sum(
            1 for weights in case_execs.values()
            for w in weights if w and w < 0
        )
        cases_rows.append([
            cid,
            c["summary"],
            c["case_status__name"] or "",
            total,
            passed,
            failed,
        ])

    return plan, runs_rows, cases_rows


PLAN_DD_RUN_HEADERS = [
    "ID", "Summary", "Started", "Stopped", "Executions", "Passed", "Completion",
]
PLAN_DD_CASE_HEADERS = [
    "ID", "Summary", "Status", "Total Executions", "Passed", "Failed",
]


@login_required
def api_consolidated_plan_report(request, plan_id):
    """CSV export for plan drill-down."""
    try:
        plan, runs_rows, cases_rows = _get_plan_drilldown_data(plan_id)
    except TestPlan.DoesNotExist:
        return JsonResponse({"error": "Test plan not found"}, status=404)

    products = [plan.product.name] if plan.product else []
    plans_meta = ["TP-{}: {}".format(plan.pk, plan.name)]
    metadata = _build_report_metadata(
        request, "Plan Drill-Down", products, plans_meta
    )

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = (
        'attachment; filename="plan_drilldown_TP{}.csv"'.format(plan_id)
    )

    writer = csv.writer(response)
    _write_csv_metadata(writer, metadata)
    writer.writerow(["Plan Drill-Down: {} (TP-{})".format(plan.name, plan.pk)])
    writer.writerow([])
    writer.writerow(["Runs"])
    writer.writerow(PLAN_DD_RUN_HEADERS)
    for row in runs_rows:
        writer.writerow(row)
    writer.writerow([])
    writer.writerow(["Cases"])
    writer.writerow(PLAN_DD_CASE_HEADERS)
    for row in cases_rows:
        writer.writerow(row)

    return response


@login_required
def api_consolidated_plan_report_excel(request, plan_id):
    """Excel export for plan drill-down."""
    from openpyxl import Workbook

    try:
        plan, runs_rows, cases_rows = _get_plan_drilldown_data(plan_id)
    except TestPlan.DoesNotExist:
        return JsonResponse({"error": "Test plan not found"}, status=404)

    products = [plan.product.name] if plan.product else []
    plans_meta = ["TP-{}: {}".format(plan.pk, plan.name)]
    metadata = _build_report_metadata(
        request, "Plan Drill-Down", products, plans_meta
    )

    wb = Workbook()

    # Report Info sheet
    ws_info = wb.active
    ws_info.title = "Report Info"
    _write_excel_metadata(ws_info, metadata)
    for col in ws_info.columns:
        max_length = 0
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws_info.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)

    ws_runs = wb.create_sheet("Runs")
    _write_excel_table(ws_runs, PLAN_DD_RUN_HEADERS, runs_rows, start_row=1)

    ws_cases = wb.create_sheet("Cases")
    _write_excel_table(ws_cases, PLAN_DD_CASE_HEADERS, cases_rows, start_row=1)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = (
        'attachment; filename="plan_drilldown_TP{}.xlsx"'.format(plan_id)
    )
    return response


@login_required
def api_consolidated_plan_report_docx(request, plan_id):
    """Word export for plan drill-down."""
    from docx import Document

    try:
        plan, runs_rows, cases_rows = _get_plan_drilldown_data(plan_id)
    except TestPlan.DoesNotExist:
        return JsonResponse({"error": "Test plan not found"}, status=404)

    products = [plan.product.name] if plan.product else []
    plans_meta = ["TP-{}: {}".format(plan.pk, plan.name)]
    metadata = _build_report_metadata(
        request, "Plan Drill-Down", products, plans_meta
    )

    doc = Document()
    doc.add_heading(
        "Plan Drill-Down: {} (TP-{})".format(plan.name, plan.pk), level=1
    )
    _write_docx_metadata(doc, metadata)

    _write_docx_table(doc, "Runs", PLAN_DD_RUN_HEADERS, runs_rows)
    _write_docx_table(doc, "Cases", PLAN_DD_CASE_HEADERS, cases_rows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    response["Content-Disposition"] = (
        'attachment; filename="plan_drilldown_TP{}.docx"'.format(plan_id)
    )
    return response


@login_required
def api_consolidated_plan_report_pdf(request, plan_id):
    """PDF export for plan drill-down."""
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table

    try:
        plan, runs_rows, cases_rows = _get_plan_drilldown_data(plan_id)
    except TestPlan.DoesNotExist:
        return JsonResponse({"error": "Test plan not found"}, status=404)

    products = [plan.product.name] if plan.product else []
    plans_meta = ["TP-{}: {}".format(plan.pk, plan.name)]
    metadata = _build_report_metadata(
        request, "Plan Drill-Down", products, plans_meta
    )

    styles = getSampleStyleSheet()
    cell_style = styles["Normal"]
    cell_style.fontSize = 7
    cell_style.leading = 9

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(letter), topMargin=0.5 * inch)
    elements = []

    elements.append(Paragraph(
        "Plan Drill-Down: {} (TP-{})".format(plan.name, plan.pk),
        styles["Title"],
    ))
    elements.append(Spacer(1, 12))
    elements.extend(_build_pdf_metadata_elements(metadata, styles))

    table_style = _get_pdf_table_style()

    # Runs table
    elements.append(Paragraph("Runs", styles["Heading2"]))
    elements.append(Spacer(1, 6))
    runs_data = [PLAN_DD_RUN_HEADERS]
    for row in runs_rows:
        r = list(row)
        r[1] = Paragraph(str(r[1]), cell_style)
        runs_data.append(r)
    run_widths = [
        0.4 * inch, 2.5 * inch, 0.8 * inch, 0.8 * inch,
        0.7 * inch, 0.6 * inch, 0.8 * inch,
    ]
    t1 = Table(runs_data, colWidths=run_widths, repeatRows=1)
    t1.setStyle(table_style)
    elements.append(t1)
    elements.append(Spacer(1, 18))

    # Cases table
    elements.append(Paragraph("Cases", styles["Heading2"]))
    elements.append(Spacer(1, 6))
    cases_data = [PLAN_DD_CASE_HEADERS]
    for row in cases_rows:
        r = list(row)
        r[1] = Paragraph(str(r[1]), cell_style)
        cases_data.append(r)
    case_widths = [
        0.5 * inch, 3.0 * inch, 0.8 * inch,
        0.8 * inch, 0.6 * inch, 0.6 * inch,
    ]
    t2 = Table(cases_data, colWidths=case_widths, repeatRows=1)
    t2.setStyle(table_style)
    elements.append(t2)

    doc.build(elements)
    buf.seek(0)

    response = HttpResponse(buf.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = (
        'attachment; filename="plan_drilldown_TP{}.pdf"'.format(plan_id)
    )
    return response


# ==========================================
# Consolidated Case Drill-Down Exports
# ==========================================


def _get_case_drilldown_data(case_id):
    """Gather plans and execution timeline data for a case drill-down export."""
    tc = TestCase.objects.select_related(
        "case_status", "priority", "author", "category", "category__product",
    ).get(pk=case_id)

    plans = list(
        tc.plan.select_related("product")
        .order_by("id")
        .values("id", "name", "product__name", "is_active")
    )
    plans_rows = [
        [
            p["id"],
            p["name"],
            p["product__name"] or "",
            "Active" if p["is_active"] else "Inactive",
        ]
        for p in plans
    ]

    executions = list(
        TestExecution.objects.filter(case=tc)
        .select_related("run", "run__plan", "status", "tested_by")
        .order_by("-run__id")
        .values(
            "run_id", "run__summary", "run__plan__name",
            "status__name", "tested_by__username", "stop_date",
        )
    )
    timeline_rows = [
        [
            e["run_id"],
            e["run__summary"],
            e["run__plan__name"] or "",
            e["status__name"] or "",
            e["tested_by__username"] or "",
            e["stop_date"].strftime("%Y-%m-%d") if e["stop_date"] else "",
        ]
        for e in executions
    ]

    return tc, plans_rows, timeline_rows


CASE_DD_PLAN_HEADERS = ["ID", "Name", "Product", "Status"]
CASE_DD_TIMELINE_HEADERS = [
    "Run ID", "Run Summary", "Plan", "Status", "Tested By", "Date",
]


@login_required
def api_consolidated_case_report(request, case_id):
    """CSV export for case drill-down."""
    try:
        tc, plans_rows, timeline_rows = _get_case_drilldown_data(case_id)
    except TestCase.DoesNotExist:
        return JsonResponse({"error": "Test case not found"}, status=404)

    products = [tc.category.product.name] if tc.category and tc.category.product else []
    plans_meta = ["TP-{}: {}".format(r[0], r[1]) for r in plans_rows]
    metadata = _build_report_metadata(
        request, "Case Drill-Down", products, plans_meta
    )

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = (
        'attachment; filename="case_drilldown_TC{}.csv"'.format(case_id)
    )

    writer = csv.writer(response)
    _write_csv_metadata(writer, metadata)
    writer.writerow(["Case Drill-Down: {} (TC-{})".format(tc.summary, tc.pk)])
    writer.writerow([])
    writer.writerow(["Plans"])
    writer.writerow(CASE_DD_PLAN_HEADERS)
    for row in plans_rows:
        writer.writerow(row)
    writer.writerow([])
    writer.writerow(["Execution Timeline"])
    writer.writerow(CASE_DD_TIMELINE_HEADERS)
    for row in timeline_rows:
        writer.writerow(row)

    return response


@login_required
def api_consolidated_case_report_excel(request, case_id):
    """Excel export for case drill-down."""
    from openpyxl import Workbook

    try:
        tc, plans_rows, timeline_rows = _get_case_drilldown_data(case_id)
    except TestCase.DoesNotExist:
        return JsonResponse({"error": "Test case not found"}, status=404)

    products = [tc.category.product.name] if tc.category and tc.category.product else []
    plans_meta = ["TP-{}: {}".format(r[0], r[1]) for r in plans_rows]
    metadata = _build_report_metadata(
        request, "Case Drill-Down", products, plans_meta
    )

    wb = Workbook()

    # Report Info sheet
    ws_info = wb.active
    ws_info.title = "Report Info"
    _write_excel_metadata(ws_info, metadata)
    for col in ws_info.columns:
        max_length = 0
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws_info.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)

    ws_plans = wb.create_sheet("Plans")
    _write_excel_table(ws_plans, CASE_DD_PLAN_HEADERS, plans_rows, start_row=1)

    ws_timeline = wb.create_sheet("Execution Timeline")
    _write_excel_table(ws_timeline, CASE_DD_TIMELINE_HEADERS, timeline_rows, start_row=1)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = (
        'attachment; filename="case_drilldown_TC{}.xlsx"'.format(case_id)
    )
    return response


@login_required
def api_consolidated_case_report_docx(request, case_id):
    """Word export for case drill-down."""
    from docx import Document

    try:
        tc, plans_rows, timeline_rows = _get_case_drilldown_data(case_id)
    except TestCase.DoesNotExist:
        return JsonResponse({"error": "Test case not found"}, status=404)

    products = [tc.category.product.name] if tc.category and tc.category.product else []
    plans_meta = ["TP-{}: {}".format(r[0], r[1]) for r in plans_rows]
    metadata = _build_report_metadata(
        request, "Case Drill-Down", products, plans_meta
    )

    doc = Document()
    doc.add_heading(
        "Case Drill-Down: {} (TC-{})".format(tc.summary, tc.pk), level=1
    )
    _write_docx_metadata(doc, metadata)

    _write_docx_table(doc, "Plans", CASE_DD_PLAN_HEADERS, plans_rows)
    _write_docx_table(doc, "Execution Timeline", CASE_DD_TIMELINE_HEADERS, timeline_rows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    response["Content-Disposition"] = (
        'attachment; filename="case_drilldown_TC{}.docx"'.format(case_id)
    )
    return response


@login_required
def api_consolidated_case_report_pdf(request, case_id):
    """PDF export for case drill-down."""
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table

    try:
        tc, plans_rows, timeline_rows = _get_case_drilldown_data(case_id)
    except TestCase.DoesNotExist:
        return JsonResponse({"error": "Test case not found"}, status=404)

    products = [tc.category.product.name] if tc.category and tc.category.product else []
    plans_meta = ["TP-{}: {}".format(r[0], r[1]) for r in plans_rows]
    metadata = _build_report_metadata(
        request, "Case Drill-Down", products, plans_meta
    )

    styles = getSampleStyleSheet()
    cell_style = styles["Normal"]
    cell_style.fontSize = 7
    cell_style.leading = 9

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(letter), topMargin=0.5 * inch)
    elements = []

    elements.append(Paragraph(
        "Case Drill-Down: {} (TC-{})".format(tc.summary, tc.pk),
        styles["Title"],
    ))
    elements.extend(_build_pdf_metadata_elements(metadata, styles))
    elements.append(Spacer(1, 12))

    table_style = _get_pdf_table_style()

    # Plans table
    elements.append(Paragraph("Plans", styles["Heading2"]))
    elements.append(Spacer(1, 6))
    plans_data = [CASE_DD_PLAN_HEADERS]
    for row in plans_rows:
        r = list(row)
        r[1] = Paragraph(str(r[1]), cell_style)
        plans_data.append(r)
    plan_widths = [
        0.5 * inch, 3.0 * inch, 1.5 * inch, 0.8 * inch,
    ]
    t1 = Table(plans_data, colWidths=plan_widths, repeatRows=1)
    t1.setStyle(table_style)
    elements.append(t1)
    elements.append(Spacer(1, 18))

    # Timeline table
    elements.append(Paragraph("Execution Timeline", styles["Heading2"]))
    elements.append(Spacer(1, 6))
    timeline_data = [CASE_DD_TIMELINE_HEADERS]
    for row in timeline_rows:
        r = list(row)
        r[1] = Paragraph(str(r[1]), cell_style)
        timeline_data.append(r)
    timeline_widths = [
        0.5 * inch, 2.5 * inch, 1.5 * inch,
        0.8 * inch, 0.8 * inch, 0.8 * inch,
    ]
    t2 = Table(timeline_data, colWidths=timeline_widths, repeatRows=1)
    t2.setStyle(table_style)
    elements.append(t2)

    doc.build(elements)
    buf.seek(0)

    response = HttpResponse(buf.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = (
        'attachment; filename="case_drilldown_TC{}.pdf"'.format(case_id)
    )
    return response
